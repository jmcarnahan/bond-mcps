"""Integration tests for the edit_word_document MCP tool."""

import io
import json
from unittest.mock import patch

import httpx
import pytest
import respx
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from ms_graph.graph_client import GRAPH_BASE_URL

from .conftest import SAMPLE_DRIVE_ITEM_WORD


def _mock_token(token: str = "test-ms-token"):
    return patch("ms_graph_mcp.get_graph_token", return_value=token)


def _get_text(result) -> str:
    return result.content[0].text


def _make_docx_bytes(*paragraphs: str) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


ITEM_ID = SAMPLE_DRIVE_ITEM_WORD["id"]
DRIVE_ID = SAMPLE_DRIVE_ITEM_WORD["parentReference"]["driveId"]

SAMPLE_UPLOADED_WORD = {
    "id": ITEM_ID,
    "name": "template.docx",
    "size": 30_000,
    "webUrl": "https://contoso.sharepoint.com/sites/engineering/template.docx",
}


@pytest.fixture
def mcp_server():
    from ms_graph_mcp import mcp

    return mcp


@pytest.fixture
def sample_doc_bytes():
    return _make_docx_bytes("Hello world", "Second paragraph", "Third paragraph")


class TestEditWordDocumentTool:
    @respx.mock
    async def test_basic_replace_edit(self, mcp_server, sample_doc_bytes):
        """Download, replace text, re-upload round-trip."""
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        upload_route = respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps([{"op": "replace", "find": "world", "replace": "earth"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        text = _get_text(result)
        assert "edited successfully" in text
        assert "Track Changes" in text
        assert upload_route.called

        # Verify the uploaded bytes contain track changes
        uploaded_bytes = upload_route.calls[0].request.content
        doc = Document(io.BytesIO(uploaded_bytes))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:del" in body_xml
        assert "w:ins" in body_xml

    @respx.mock
    async def test_unsupported_file_rejected(self, mcp_server):
        non_doc_item = {**SAMPLE_DRIVE_ITEM_WORD, "name": "data.txt"}
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=non_doc_item)
        )

        edits = json.dumps([{"op": "append", "content": "new"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "not an editable document" in _get_text(result)

    @respx.mock
    async def test_invalid_edits_json(self, mcp_server):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": "not json{",
                    },
                )

        assert "Invalid edits" in _get_text(result)

    @respx.mock
    async def test_text_not_found_error(self, mcp_server, sample_doc_bytes):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )

        edits = json.dumps([{"op": "replace", "find": "nonexistent", "replace": "x"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "Edit failed" in _get_text(result)
        assert "Text not found" in _get_text(result)

    @respx.mock
    async def test_track_changes_disabled(self, mcp_server, sample_doc_bytes):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        upload_route = respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps([{"op": "replace", "find": "world", "replace": "earth"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                        "options": json.dumps({"track_changes": False}),
                    },
                )

        text = _get_text(result)
        assert "edited successfully" in text
        assert "Track Changes" not in text

        # Verify no revision markup in uploaded doc
        uploaded_bytes = upload_route.calls[0].request.content
        doc = Document(io.BytesIO(uploaded_bytes))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:del" not in body_xml

    @respx.mock
    async def test_custom_author(self, mcp_server, sample_doc_bytes):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        upload_route = respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps([{"op": "replace", "find": "world", "replace": "earth"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                        "options": json.dumps({"author": "Aitor González"}),
                    },
                )

        text = _get_text(result)
        assert "Aitor González" in text

        # Verify author in the revision markup
        uploaded_bytes = upload_route.calls[0].request.content
        doc = Document(io.BytesIO(uploaded_bytes))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "Aitor" in body_xml

    @respx.mock
    async def test_comment_operation_end_to_end(self, mcp_server, sample_doc_bytes):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        upload_route = respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps(
            [{"op": "comment", "find": "Second paragraph", "comment": "Please review this"}]
        )

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "edited successfully" in _get_text(result)

        # Verify comment markers in uploaded doc
        uploaded_bytes = upload_route.calls[0].request.content
        doc = Document(io.BytesIO(uploaded_bytes))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "commentRangeStart" in body_xml

    @respx.mock
    async def test_multiple_operations(self, mcp_server, sample_doc_bytes):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps(
            [
                {"op": "replace", "find": "world", "replace": "earth"},
                {"op": "append", "content": "New final paragraph"},
                {"op": "comment", "find": "Third paragraph", "comment": "Check this"},
            ]
        )

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        text = _get_text(result)
        assert "3" in text  # 3 operations
        assert "replace, append, comment" in text

    @respx.mock
    async def test_sharepoint_site_id(self, mcp_server, sample_doc_bytes):
        site_id = "site-123"
        respx.get(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        respx.put(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps([{"op": "append", "content": "New paragraph"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                        "site_id": site_id,
                    },
                )

        assert "edited successfully" in _get_text(result)

    @respx.mock
    async def test_file_not_found(self, mcp_server):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(404, json={"error": {"code": "itemNotFound"}})
        )

        edits = json.dumps([{"op": "append", "content": "text"}])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "not found" in _get_text(result).lower() or "Error" in _get_text(result)

    @respx.mock
    async def test_empty_edits_array(self, mcp_server):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        edits = json.dumps([])

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "No edit operations" in _get_text(result)

    @respx.mock
    async def test_full_edit_workflow_produces_readable_document(
        self, mcp_server, sample_doc_bytes
    ):
        """End-to-end: edit + comment, verify the uploaded doc has correct content."""
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, content=sample_doc_bytes)
        )
        upload_route = respx.put(f"{GRAPH_BASE_URL}/me/drive/items/{ITEM_ID}/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_WORD)
        )

        edits = json.dumps(
            [
                {"op": "replace", "find": "Hello world", "replace": "Hello earth"},
                {"op": "append", "content": "Final thoughts here."},
                {"op": "comment", "find": "Second paragraph", "comment": "Please verify this."},
            ]
        )

        with _mock_token():
            from fastmcp import Client

            async with Client(mcp_server) as client:
                result = await client.call_tool(
                    "edit_document",
                    {
                        "item_id": ITEM_ID,
                        "edits": edits,
                    },
                )

        assert "edited successfully" in _get_text(result)

        # Deeply verify the uploaded document
        uploaded_bytes = upload_route.calls[0].request.content
        doc = Document(io.BytesIO(uploaded_bytes))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")

        # Track changes: replacement text present in ins, original in del
        assert "Hello earth" in body_xml
        assert "Hello world" in body_xml  # preserved in delText

        # Appended paragraph visible
        assert "Final thoughts here." in body_xml

        # Comment markers present
        assert "commentRangeStart" in body_xml

        # Verify the comments part has our comment text
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        for rel in doc.part.rels.values():
            if rel.reltype == RT.COMMENTS:
                comments_xml = etree.fromstring(rel.target_part.blob)
                t_elems = comments_xml.findall(f".//{qn('w:t')}")
                comment_texts = [t.text for t in t_elems if t.text]
                assert "Please verify this." in comment_texts
                break
        else:
            pytest.fail("No comments part found in uploaded document")
