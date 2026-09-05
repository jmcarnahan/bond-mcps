"""Tests for file/drive operations (sync and async)."""

import base64
import io
import json
from unittest.mock import patch

import httpx
import pytest
import respx
from ms_graph import files
from ms_graph.graph_client import GRAPH_BASE_URL, AsyncGraphClient, GraphClient, GraphError

from .conftest import (
    GRAPH_ERROR_403,
    GRAPH_ERROR_404,
    SAMPLE_COPY_COMPLETED,
    SAMPLE_COPY_FAILED,
    SAMPLE_COPY_IN_PROGRESS,
    SAMPLE_DRIVE_CHILDREN_RESPONSE,
    SAMPLE_DRIVE_ITEM_BINARY,
    SAMPLE_DRIVE_ITEM_FILE,
    SAMPLE_DRIVE_ITEM_FOLDER,
    SAMPLE_DRIVE_ITEM_LARGE_TEXT,
    SAMPLE_DRIVE_ITEM_WORD,
    SAMPLE_DRIVE_UPLOAD_SESSION,
    SAMPLE_DRIVE_UPLOAD_URL,
    SAMPLE_SEARCH_RESPONSE,
    SAMPLE_SEARCH_RESPONSE_EMPTY,
    SAMPLE_SHARED_DRIVE_ITEM,
    SAMPLE_SHARED_FOLDER_CHILDREN,
    SAMPLE_SHARED_TEXT_FILE,
    SAMPLE_SITE,
    SAMPLE_SITES_RESPONSE,
    SAMPLE_UPLOADED_FILE,
)


class TestFilesSync:
    """Synchronous file operation tests."""

    @respx.mock
    def test_list_drive_children_root(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_CHILDREN_RESPONSE)
        )
        with GraphClient("tok") as client:
            items = files.list_drive_children(client)

        assert len(items) == 3
        assert items[0]["name"] == "Documents"
        assert items[1]["name"] == "report.csv"

    @respx.mock
    def test_list_drive_children_subfolder(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root:/Documents:/children").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        with GraphClient("tok") as client:
            items = files.list_drive_children(client, folder_path="Documents")

        assert len(items) == 1
        assert items[0]["name"] == "report.csv"

    @respx.mock
    def test_list_drive_children_sharepoint(self):
        respx.get(f"{GRAPH_BASE_URL}/sites/site-id-001/drive/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_CHILDREN_RESPONSE)
        )
        with GraphClient("tok") as client:
            items = files.list_drive_children(client, site_id="site-id-001")

        assert len(items) == 3

    @respx.mock
    def test_get_drive_item(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_FILE)
        )
        with GraphClient("tok") as client:
            item = files.get_drive_item(client, "file-id-001")

        assert item["name"] == "report.csv"
        assert item["size"] == 1024

    @respx.mock
    def test_get_drive_item_content_text(self):
        csv_content = b"header1,header2\nvalue1,value2\n"
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_FILE)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001/content").mock(
            return_value=httpx.Response(200, content=csv_content)
        )
        with GraphClient("tok") as client:
            item, content = files.get_drive_item_content(client, "file-id-001")

        assert item["name"] == "report.csv"
        assert content == "header1,header2\nvalue1,value2\n"

    @respx.mock
    def test_get_drive_item_content_binary(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-002").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_BINARY)
        )
        with GraphClient("tok") as client:
            item, content = files.get_drive_item_content(client, "file-id-002")

        assert item["name"] == "presentation.pptx"
        assert content is None

    @respx.mock
    def test_get_drive_item_content_too_large(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-003").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_LARGE_TEXT)
        )
        with GraphClient("tok") as client:
            item, content = files.get_drive_item_content(client, "file-id-003")

        assert item["name"] == "huge-log.txt"
        assert content is None  # Too large, skipped

    @respx.mock
    def test_search_drive(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/search(q='report')").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        with GraphClient("tok") as client:
            results = files.search_drive(client, "report")

        assert len(results) == 1
        assert results[0]["name"] == "report.csv"

    @respx.mock
    def test_search_files_unified(self):
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(200, json=SAMPLE_SEARCH_RESPONSE)
        )
        with GraphClient("tok") as client:
            results = files.search_files_unified(client, "budget")

        assert len(results) == 2
        assert results[0]["name"] == "Q4-budget.xlsx"
        assert results[0]["_searchSummary"] == "Q4 <c0>budget</c0> projections for 2025"
        assert results[1]["name"] == "budget-notes.md"

    @respx.mock
    def test_search_files_unified_empty(self):
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(200, json=SAMPLE_SEARCH_RESPONSE_EMPTY)
        )
        with GraphClient("tok") as client:
            results = files.search_files_unified(client, "nonexistent")

        assert results == []

    @respx.mock
    def test_list_sites(self):
        respx.get(f"{GRAPH_BASE_URL}/sites").mock(
            return_value=httpx.Response(200, json=SAMPLE_SITES_RESPONSE)
        )
        with GraphClient("tok") as client:
            sites = files.list_sites(client, query="engineering")

        assert len(sites) == 2
        assert sites[0]["displayName"] == "Engineering Hub"

    @respx.mock
    def test_search_files_unified_consumer_fallback(self):
        """Consumer accounts get 400 from /search/query — should fall back to per-drive search."""
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(
                400,
                json={
                    "error": {
                        "code": "BadRequest",
                        "message": "This API is not supported for MSA accounts",
                    }
                },
            )
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/search(q='report')").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        with GraphClient("tok") as client:
            results = files.search_files_unified(client, "report")

        assert len(results) == 1
        assert results[0]["name"] == "report.csv"

    @respx.mock
    def test_list_sites_followed(self):
        respx.get(f"{GRAPH_BASE_URL}/me/followedSites").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_SITE]})
        )
        with GraphClient("tok") as client:
            sites = files.list_sites(client)

        assert len(sites) == 1
        assert sites[0]["displayName"] == "Engineering Hub"


class TestFilesAsync:
    """Async file operation tests."""

    @respx.mock
    async def test_alist_drive_children_root(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_CHILDREN_RESPONSE)
        )
        async with AsyncGraphClient("tok") as client:
            items = await files.alist_drive_children(client)

        assert len(items) == 3

    @respx.mock
    async def test_alist_drive_children_subfolder(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root:/Projects:/children").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        async with AsyncGraphClient("tok") as client:
            items = await files.alist_drive_children(client, folder_path="Projects")

        assert len(items) == 1

    @respx.mock
    async def test_alist_drive_children_sharepoint(self):
        respx.get(f"{GRAPH_BASE_URL}/sites/site-id-001/drive/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_CHILDREN_RESPONSE)
        )
        async with AsyncGraphClient("tok") as client:
            items = await files.alist_drive_children(client, site_id="site-id-001")

        assert len(items) == 3

    @respx.mock
    async def test_aget_drive_item(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.aget_drive_item(client, "file-id-001")

        assert item["name"] == "report.csv"

    @respx.mock
    async def test_aget_drive_item_content_text(self):
        csv_content = b"col1,col2\na,b\n"
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_FILE)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-001/content").mock(
            return_value=httpx.Response(200, content=csv_content)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_content(client, "file-id-001")

        assert content == "col1,col2\na,b\n"

    @respx.mock
    async def test_aget_drive_item_content_binary(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-002").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_BINARY)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_content(client, "file-id-002")

        assert content is None

    @respx.mock
    async def test_aget_drive_item_content_too_large(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-003").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_LARGE_TEXT)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_content(client, "file-id-003")

        assert content is None

    @respx.mock
    async def test_asearch_drive(self):
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/search(q='report')").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        async with AsyncGraphClient("tok") as client:
            results = await files.asearch_drive(client, "report")

        assert len(results) == 1

    @respx.mock
    async def test_asearch_files_unified(self):
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(200, json=SAMPLE_SEARCH_RESPONSE)
        )
        async with AsyncGraphClient("tok") as client:
            results = await files.asearch_files_unified(client, "budget")

        assert len(results) == 2
        assert "_searchSummary" in results[0]

    @respx.mock
    async def test_asearch_files_unified_empty(self):
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(200, json=SAMPLE_SEARCH_RESPONSE_EMPTY)
        )
        async with AsyncGraphClient("tok") as client:
            results = await files.asearch_files_unified(client, "nonexistent")

        assert results == []

    @respx.mock
    async def test_alist_sites(self):
        respx.get(f"{GRAPH_BASE_URL}/sites").mock(
            return_value=httpx.Response(200, json=SAMPLE_SITES_RESPONSE)
        )
        async with AsyncGraphClient("tok") as client:
            sites = await files.alist_sites(client, query="engineering")

        assert len(sites) == 2

    @respx.mock
    async def test_asearch_files_unified_consumer_fallback(self):
        """Consumer accounts get 400 from /search/query — should fall back to per-drive search."""
        respx.post(f"{GRAPH_BASE_URL}/search/query").mock(
            return_value=httpx.Response(
                400,
                json={
                    "error": {
                        "code": "BadRequest",
                        "message": "This API is not supported for MSA accounts",
                    }
                },
            )
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/root/search(q='report')").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_DRIVE_ITEM_FILE]})
        )
        async with AsyncGraphClient("tok") as client:
            results = await files.asearch_files_unified(client, "report")

        assert len(results) == 1
        assert results[0]["name"] == "report.csv"

    @respx.mock
    async def test_alist_sites_followed(self):
        respx.get(f"{GRAPH_BASE_URL}/me/followedSites").mock(
            return_value=httpx.Response(200, json={"value": [SAMPLE_SITE]})
        )
        async with AsyncGraphClient("tok") as client:
            sites = await files.alist_sites(client)

        assert len(sites) == 1


# ---------------------------------------------------------------------------
# Upload tests
# ---------------------------------------------------------------------------


class TestUploadSync:
    """Synchronous upload_file tests."""

    @respx.mock
    def test_upload_to_root(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/report.md:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            item = files.upload_file(
                client, folder_path="", filename="report.md", content="# Hello"
            )

        assert item["name"] == "report.md"
        assert route.calls[0].request.headers["Content-Type"] == "text/markdown"
        assert route.calls[0].request.content == b"# Hello"

    @respx.mock
    def test_upload_to_subfolder(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/Documents/data.csv:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            files.upload_file(
                client, folder_path="Documents", filename="data.csv", content="a,b\n1,2"
            )

        assert route.called
        assert route.calls[0].request.headers["Content-Type"] == "text/csv"

    @respx.mock
    def test_upload_to_sharepoint(self):
        site_id = "site-id-001"
        route = respx.put(
            f"{GRAPH_BASE_URL}/sites/{site_id}/drive/root:/Shared Documents/notes.txt:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        with GraphClient("tok") as client:
            files.upload_file(
                client,
                folder_path="Shared Documents",
                filename="notes.txt",
                content="hello",
                site_id=site_id,
            )
        assert route.called
        assert route.calls[0].request.headers["Content-Type"] == "text/plain"

    @respx.mock
    def test_upload_overwrites_existing(self):
        """200 response means file was updated (overwritten)."""
        respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/existing.json:/content").mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            item = files.upload_file(client, folder_path="", filename="existing.json", content="{}")
        assert item is not None

    @respx.mock
    def test_upload_unknown_extension_uses_octet_stream(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/file.bin:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            files.upload_file(client, folder_path="", filename="file.bin", content="data")
        assert route.calls[0].request.headers["Content-Type"] == "application/octet-stream"

    @pytest.mark.parametrize(
        "ext,expected_ct",
        [
            ("report.txt", "text/plain"),
            ("page.html", "text/html"),
            ("data.json", "application/json"),
            ("config.yaml", "application/yaml"),
            ("config.yml", "application/yaml"),
            ("schema.xml", "application/xml"),
        ],
    )
    @respx.mock
    def test_upload_content_type_inference(self, ext, expected_ct):
        route = respx.put(url__regex=r"/content$").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            files.upload_file(client, folder_path="", filename=ext, content="x")
        assert route.called
        ct = route.calls[0].request.headers["Content-Type"]
        assert ct == expected_ct

    def test_upload_rejects_content_over_4mb(self):
        """Content exceeding 4 MB raises ValueError before making an API call."""
        big_content = "x" * 4_000_001
        with pytest.raises(ValueError, match="4 MB"):
            with GraphClient("tok") as client:
                files.upload_file(client, folder_path="", filename="big.txt", content=big_content)

    async def test_aupload_rejects_content_over_4mb(self):
        """Async upload also rejects content over 4 MB."""
        big_content = "x" * 4_000_001
        with pytest.raises(ValueError, match="4 MB"):
            async with AsyncGraphClient("tok") as client:
                await files.aupload_file(
                    client, folder_path="", filename="big.txt", content=big_content
                )


class TestUploadAsync:
    """Asynchronous aupload_file tests."""

    @respx.mock
    async def test_aupload_to_root(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/notes.md:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.aupload_file(
                client, folder_path="", filename="notes.md", content="hello"
            )
        assert item["id"] == SAMPLE_UPLOADED_FILE["id"]
        assert route.calls[0].request.headers["Content-Type"] == "text/markdown"

    @respx.mock
    async def test_aupload_to_subfolder(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/Reports/summary.csv:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            await files.aupload_file(
                client, folder_path="Reports", filename="summary.csv", content="x,y"
            )
        assert route.called
        assert route.calls[0].request.headers["Content-Type"] == "text/csv"

    @respx.mock
    async def test_aupload_to_sharepoint(self):
        site_id = "site-id-001"
        route = respx.put(
            f"{GRAPH_BASE_URL}/sites/{site_id}/drive/root:/Docs/page.html:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        async with AsyncGraphClient("tok") as client:
            await files.aupload_file(
                client,
                folder_path="Docs",
                filename="page.html",
                content="<p>hi</p>",
                site_id=site_id,
            )
        assert route.called

    @respx.mock
    async def test_aupload_bytes_pdf(self):
        """aupload_bytes uploads raw binary with the specified content-type."""
        pdf_bytes = b"%PDF-1.4 fake content"
        route = respx.put(
            f"{GRAPH_BASE_URL}/me/drive/root:/Power BI Exports/report.pdf:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        async with AsyncGraphClient("tok") as client:
            item = await files.aupload_bytes(
                client,
                folder_path="Power BI Exports",
                filename="report.pdf",
                data=pdf_bytes,
                content_type="application/pdf",
            )

        assert item["id"] == SAMPLE_UPLOADED_FILE["id"]
        assert route.calls[0].request.headers["Content-Type"] == "application/pdf"
        assert route.calls[0].request.content == pdf_bytes

    async def test_aupload_bytes_rejects_over_4mb(self):
        with pytest.raises(ValueError, match="4 MB"):
            async with AsyncGraphClient("tok") as client:
                await files.aupload_bytes(
                    client,
                    folder_path="",
                    filename="big.pdf",
                    data=b"x" * 4_000_001,
                    content_type="application/pdf",
                )


# ---------------------------------------------------------------------------
# Copy tests
# ---------------------------------------------------------------------------

MONITOR_URL = "https://api.onedrive.com/v1.0/monitor/copy-op-token"
SOURCE_DRIVE_ID = SAMPLE_DRIVE_ITEM_WORD["parentReference"]["driveId"]


class TestCopySync:
    """Synchronous copy_drive_item tests."""

    @pytest.fixture(autouse=True)
    def patch_sleep(self, no_sleep):
        pass

    @respx.mock
    def test_copy_succeeds_after_one_poll_200(self):
        """200 monitor response with completed status."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            result = files.copy_drive_item(client, item_id=item_id, new_name="template-copy.docx")

        assert result["status"] == "completed"
        assert result["resourceId"] == SAMPLE_COPY_COMPLETED["resourceId"]

    @respx.mock
    def test_copy_succeeds_via_303(self):
        """303 monitor response (real SharePoint behavior) — completed status in body."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(303, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            result = files.copy_drive_item(client, item_id=item_id, new_name="template-copy.docx")

        assert result["status"] == "completed"
        assert result["resourceId"] == SAMPLE_COPY_COMPLETED["resourceId"]

    @respx.mock
    def test_copy_polls_until_completed(self):
        """In-progress response is polled past before completed fires."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        # First poll → inProgress (200), second poll → completed (303)
        respx.get(MONITOR_URL).mock(
            side_effect=[
                httpx.Response(200, json=SAMPLE_COPY_IN_PROGRESS),
                httpx.Response(303, json=SAMPLE_COPY_COMPLETED),
            ]
        )
        with GraphClient("tok") as client:
            result = files.copy_drive_item(client, item_id=item_id, new_name="copy.docx")

        assert result["status"] == "completed"
        assert respx.calls.call_count == 4  # GET item + POST copy + 2x GET monitor

    @respx.mock
    def test_copy_raises_on_failure(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_FAILED))
        with pytest.raises(GraphError, match="accessDenied"):
            with GraphClient("tok") as client:
                files.copy_drive_item(client, item_id=item_id, new_name="copy.docx")

    @respx.mock
    def test_copy_with_explicit_destination_folder(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        dest_folder_id = "folder-id-archive"
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        copy_route = respx.post(
            f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy"
        ).mock(return_value=httpx.Response(202, headers={"Location": MONITOR_URL}))
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            files.copy_drive_item(
                client,
                item_id=item_id,
                new_name="archived.docx",
                destination_folder_id=dest_folder_id,
            )

        copy_body = json.loads(copy_route.calls[0].request.content)
        assert copy_body["parentReference"]["id"] == dest_folder_id

    @respx.mock
    def test_copy_on_sharepoint(self):
        site_id = "site-id-001"
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            result = files.copy_drive_item(
                client, item_id=item_id, new_name="sp-copy.docx", site_id=site_id
            )
        assert result["status"] == "completed"

    @respx.mock
    def test_copy_falls_back_when_no_source_drive_id(self):
        """Falls back to site-based path when source has no driveId in parentReference."""
        site_id = "site-id-fallback"
        source_no_drive = {
            **SAMPLE_DRIVE_ITEM_WORD,
            "parentReference": {"id": "folder-id-root", "path": "/drive/root:"},
        }
        item_id = source_no_drive["id"]
        respx.get(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=source_no_drive)
        )
        respx.post(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            result = files.copy_drive_item(
                client, item_id=item_id, new_name="fallback-copy.docx", site_id=site_id
            )
        assert result["status"] == "completed"

    @respx.mock
    def test_copy_passes_drive_id_from_source(self):
        """driveId from source parentReference is included in the copy request."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        source_drive_id = SAMPLE_DRIVE_ITEM_WORD["parentReference"]["driveId"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        copy_route = respx.post(
            f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy"
        ).mock(return_value=httpx.Response(202, headers={"Location": MONITOR_URL}))
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            files.copy_drive_item(client, item_id=item_id, new_name="copy.docx")

        copy_body = json.loads(copy_route.calls[0].request.content)
        assert copy_body["parentReference"]["driveId"] == source_drive_id
        assert copy_body["name"] == "copy.docx"

    @respx.mock
    def test_copy_raises_on_missing_location_header(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202)  # No Location header
        )
        with pytest.raises(GraphError, match="NoLocation"):
            with GraphClient("tok") as client:
                files.copy_drive_item(client, item_id=item_id, new_name="copy.docx")

    @respx.mock
    def test_copy_times_out(self):
        """Raises CopyTimeout when the operation never completes within the deadline."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_IN_PROGRESS))
        import ms_graph.files as files_mod

        with patch.object(files_mod, "_COPY_POLL_TIMEOUT", 0):
            with pytest.raises(GraphError, match="CopyTimeout"):
                with GraphClient("tok") as client:
                    files.copy_drive_item(client, item_id=item_id, new_name="copy.docx")

    @respx.mock
    def test_copy_with_explicit_source_drive_id(self):
        """source_drive_id routes the initial GET to /drives/{drive_id}/items/..."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        custom_drive = "drive-custom-source-001"
        get_route = respx.get(f"{GRAPH_BASE_URL}/drives/{custom_drive}/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        with GraphClient("tok") as client:
            result = files.copy_drive_item(
                client,
                item_id=item_id,
                new_name="source-drive-copy.docx",
                source_drive_id=custom_drive,
            )

        assert get_route.called
        assert result["status"] == "completed"


class TestCopyAsync:
    """Asynchronous acopy_drive_item tests."""

    @pytest.fixture(autouse=True)
    def patch_sleep(self, no_sleep):
        pass

    @respx.mock
    async def test_acopy_succeeds_200(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        async with AsyncGraphClient("tok") as client:
            result = await files.acopy_drive_item(
                client, item_id=item_id, new_name="async-copy.docx"
            )

        assert result["status"] == "completed"
        assert result["resourceId"] == SAMPLE_COPY_COMPLETED["resourceId"]

    @respx.mock
    async def test_acopy_succeeds_303(self):
        """303 monitor response (real SharePoint behavior) — completed status in body."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(303, json=SAMPLE_COPY_COMPLETED))
        async with AsyncGraphClient("tok") as client:
            result = await files.acopy_drive_item(
                client, item_id=item_id, new_name="async-copy.docx"
            )

        assert result["status"] == "completed"
        assert result["resourceId"] == SAMPLE_COPY_COMPLETED["resourceId"]

    @respx.mock
    async def test_acopy_raises_on_failure(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_FAILED))
        with pytest.raises(GraphError, match="accessDenied"):
            async with AsyncGraphClient("tok") as client:
                await files.acopy_drive_item(client, item_id=item_id, new_name="copy.docx")

    @respx.mock
    async def test_acopy_with_explicit_destination(self):
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        dest = "folder-id-target"
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        copy_route = respx.post(
            f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy"
        ).mock(return_value=httpx.Response(202, headers={"Location": MONITOR_URL}))
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        async with AsyncGraphClient("tok") as client:
            await files.acopy_drive_item(
                client, item_id=item_id, new_name="copy.docx", destination_folder_id=dest
            )

        copy_body = json.loads(copy_route.calls[0].request.content)
        assert copy_body["parentReference"]["id"] == dest

    @respx.mock
    async def test_acopy_times_out(self):
        """Async copy raises CopyTimeout when operation never completes."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_IN_PROGRESS))
        import ms_graph.files as files_mod

        with patch.object(files_mod, "_COPY_POLL_TIMEOUT", 0):
            with pytest.raises(GraphError, match="CopyTimeout"):
                async with AsyncGraphClient("tok") as client:
                    await files.acopy_drive_item(client, item_id=item_id, new_name="copy.docx")

    @respx.mock
    async def test_acopy_with_explicit_source_drive_id(self):
        """source_drive_id routes the initial GET to /drives/{drive_id}/items/... (async)."""
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        custom_drive = "drive-custom-source-001"
        get_route = respx.get(f"{GRAPH_BASE_URL}/drives/{custom_drive}/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_WORD)
        )
        respx.post(f"{GRAPH_BASE_URL}/drives/{SOURCE_DRIVE_ID}/items/{item_id}/copy").mock(
            return_value=httpx.Response(202, headers={"Location": MONITOR_URL})
        )
        respx.get(MONITOR_URL).mock(return_value=httpx.Response(200, json=SAMPLE_COPY_COMPLETED))
        async with AsyncGraphClient("tok") as client:
            result = await files.acopy_drive_item(
                client,
                item_id=item_id,
                new_name="source-drive-copy.docx",
                source_drive_id=custom_drive,
            )

        assert get_route.called
        assert result["status"] == "completed"


# ---------------------------------------------------------------------------
# Rename tests
# ---------------------------------------------------------------------------

SAMPLE_RENAMED_FILE = {**SAMPLE_DRIVE_ITEM_FILE, "name": "renamed-report.csv"}
SAMPLE_RENAMED_FOLDER = {**SAMPLE_DRIVE_ITEM_FOLDER, "name": "Archive-2025"}


class TestRenameSync:
    """Synchronous rename_drive_item tests."""

    @respx.mock
    def test_rename_file(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        route = respx.patch(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_RENAMED_FILE)
        )
        with GraphClient("tok") as client:
            item = files.rename_drive_item(client, item_id=item_id, new_name="renamed-report.csv")

        assert item["name"] == "renamed-report.csv"
        body = json.loads(route.calls[0].request.content)
        assert body == {"name": "renamed-report.csv"}

    @respx.mock
    def test_rename_folder(self):
        item_id = SAMPLE_DRIVE_ITEM_FOLDER["id"]
        respx.patch(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_RENAMED_FOLDER)
        )
        with GraphClient("tok") as client:
            item = files.rename_drive_item(client, item_id=item_id, new_name="Archive-2025")

        assert item["name"] == "Archive-2025"

    @respx.mock
    def test_rename_on_sharepoint(self):
        site_id = "site-id-001"
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        route = respx.patch(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_RENAMED_FILE)
        )
        with GraphClient("tok") as client:
            item = files.rename_drive_item(
                client, item_id=item_id, new_name="renamed-report.csv", site_id=site_id
            )

        assert item["name"] == "renamed-report.csv"
        assert route.called

    @respx.mock
    def test_rename_propagates_graph_error(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        respx.patch(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(
                404, json={"error": {"code": "ResourceNotFound", "message": "Item not found."}}
            )
        )
        with pytest.raises(GraphError, match="ResourceNotFound"):
            with GraphClient("tok") as client:
                files.rename_drive_item(client, item_id=item_id, new_name="new.csv")


class TestRenameAsync:
    """Asynchronous arename_drive_item tests."""

    @respx.mock
    async def test_arename_file(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        route = respx.patch(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(200, json=SAMPLE_RENAMED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.arename_drive_item(
                client, item_id=item_id, new_name="renamed-report.csv"
            )

        assert item["name"] == "renamed-report.csv"
        body = json.loads(route.calls[0].request.content)
        assert body == {"name": "renamed-report.csv"}

    @respx.mock
    async def test_arename_on_sharepoint(self):
        site_id = "site-id-001"
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        route = respx.patch(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}").mock(
            return_value=httpx.Response(
                200, json={**SAMPLE_DRIVE_ITEM_WORD, "name": "final-doc.docx"}
            )
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.arename_drive_item(
                client, item_id=item_id, new_name="final-doc.docx", site_id=site_id
            )

        assert item["name"] == "final-doc.docx"
        assert route.called

    @respx.mock
    async def test_arename_propagates_graph_error(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        respx.patch(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(
                403, json={"error": {"code": "AccessDenied", "message": "Cannot rename."}}
            )
        )
        with pytest.raises(GraphError, match="AccessDenied"):
            async with AsyncGraphClient("tok") as client:
                await files.arename_drive_item(client, item_id=item_id, new_name="x.csv")


class TestDeleteAsync:
    """Asynchronous adelete_drive_item tests."""

    @respx.mock
    async def test_adelete_file(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        route = respx.delete(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(204)
        )
        async with AsyncGraphClient("tok") as client:
            result = await files.adelete_drive_item(client, item_id=item_id)

        assert result is None
        assert route.called

    @respx.mock
    async def test_adelete_on_sharepoint(self):
        site_id = "site-id-001"
        item_id = SAMPLE_DRIVE_ITEM_WORD["id"]
        route = respx.delete(f"{GRAPH_BASE_URL}/sites/{site_id}/drive/items/{item_id}").mock(
            return_value=httpx.Response(204)
        )
        async with AsyncGraphClient("tok") as client:
            await files.adelete_drive_item(client, item_id=item_id, site_id=site_id)

        assert route.called

    @respx.mock
    async def test_adelete_propagates_graph_error(self):
        item_id = SAMPLE_DRIVE_ITEM_FILE["id"]
        respx.delete(f"{GRAPH_BASE_URL}/me/drive/items/{item_id}").mock(
            return_value=httpx.Response(
                404, json={"error": {"code": "itemNotFound", "message": "Not found."}}
            )
        )
        with pytest.raises(GraphError, match="itemNotFound"):
            async with AsyncGraphClient("tok") as client:
                await files.adelete_drive_item(client, item_id=item_id)


# ---------------------------------------------------------------------------
# Sharing URL resolution tests
# ---------------------------------------------------------------------------

SAMPLE_SHARING_URL = "https://mcafee-my.sharepoint.com/:p:/p/sajith_pilakkavil/IQDlH5omr1bEQpRD3GJa7fmqAemTOa6IJ3XnNMgxgAQZPsk"


class TestEncodeSharingUrl:
    """Unit tests for URL detection and encoding helpers."""

    def test_encode_produces_u_bang_prefix(self):
        token = files._encode_sharing_url(SAMPLE_SHARING_URL)
        assert token.startswith("u!")

    def test_encode_no_forbidden_chars(self):
        token = files._encode_sharing_url(SAMPLE_SHARING_URL)
        payload = token[2:]
        assert "/" not in payload
        assert "+" not in payload
        assert not payload.endswith("=")

    def test_encode_roundtrip(self):
        url = "https://contoso.sharepoint.com/:w:/s/site/EaB123-xyz"
        token = files._encode_sharing_url(url)
        payload = token[2:].replace("_", "/").replace("-", "+")
        padding = 4 - len(payload) % 4
        if padding != 4:
            payload += "=" * padding
        decoded = base64.b64decode(payload).decode("utf-8")
        assert decoded == url

    def test_is_sharing_url_sharepoint(self):
        assert files.is_sharing_url("https://contoso-my.sharepoint.com/:p:/p/user/abc")
        assert files.is_sharing_url("https://contoso.sharepoint.com/:w:/s/site/abc")

    def test_is_sharing_url_onedrive(self):
        assert files.is_sharing_url("https://1drv.ms/w/s!abc123")
        assert files.is_sharing_url("https://onedrive.live.com/redir?resid=abc")

    def test_is_sharing_url_negative(self):
        assert not files.is_sharing_url("file-id-001")
        assert not files.is_sharing_url("01ABCDEF12345")
        assert not files.is_sharing_url("https://google.com/doc/123")
        assert not files.is_sharing_url("")

    def test_is_sharing_url_with_whitespace(self):
        assert files.is_sharing_url("  https://contoso.sharepoint.com/:w:/s/site/abc  ")


class TestResolveSharingLinkSync:
    """Synchronous sharing link resolution tests."""

    @respx.mock
    def test_resolve_succeeds(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_DRIVE_ITEM)
        )
        with GraphClient("tok") as client:
            item = files.resolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert item["name"] == "Q4-Presentation.pptx"
        assert item["id"] == "shared-file-001"

    @respx.mock
    def test_resolve_access_denied(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(403, json=GRAPH_ERROR_403)
        )
        with pytest.raises(GraphError) as exc_info:
            with GraphClient("tok") as client:
                files.resolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert exc_info.value.status_code == 403

    @respx.mock
    def test_resolve_not_found(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(404, json=GRAPH_ERROR_404)
        )
        with pytest.raises(GraphError) as exc_info:
            with GraphClient("tok") as client:
                files.resolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert exc_info.value.status_code == 404

    @respx.mock
    def test_resolve_content_text_file(self):
        md_content = b"# Hello\n\nThis is a shared note."
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_TEXT_FILE)
        )
        respx.get(url__regex=r"/shares/u!.*/driveItem/content$").mock(
            return_value=httpx.Response(200, content=md_content)
        )
        with GraphClient("tok") as client:
            item, content = files.resolve_sharing_link_content(client, SAMPLE_SHARING_URL)
        assert item["name"] == "notes.md"
        assert content == "# Hello\n\nThis is a shared note."

    @respx.mock
    def test_resolve_content_binary_file(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_DRIVE_ITEM)
        )
        with GraphClient("tok") as client:
            item, content = files.resolve_sharing_link_content(client, SAMPLE_SHARING_URL)
        assert item["name"] == "Q4-Presentation.pptx"
        assert content is None

    @respx.mock
    def test_list_sharing_link_children(self):
        respx.get(url__regex=r"/shares/u!.*/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_FOLDER_CHILDREN)
        )
        with GraphClient("tok") as client:
            items = files.list_sharing_link_children(client, SAMPLE_SHARING_URL)
        assert len(items) == 2
        assert items[0]["name"] == "file1.docx"
        assert items[1]["name"] == "data.csv"


class TestResolveSharingLinkAsync:
    """Asynchronous sharing link resolution tests."""

    @respx.mock
    async def test_aresolve_succeeds(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_DRIVE_ITEM)
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.aresolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert item["name"] == "Q4-Presentation.pptx"
        assert item["id"] == "shared-file-001"

    @respx.mock
    async def test_aresolve_access_denied(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(403, json=GRAPH_ERROR_403)
        )
        with pytest.raises(GraphError) as exc_info:
            async with AsyncGraphClient("tok") as client:
                await files.aresolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert exc_info.value.status_code == 403

    @respx.mock
    async def test_aresolve_not_found(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(404, json=GRAPH_ERROR_404)
        )
        with pytest.raises(GraphError) as exc_info:
            async with AsyncGraphClient("tok") as client:
                await files.aresolve_sharing_link(client, SAMPLE_SHARING_URL)
        assert exc_info.value.status_code == 404

    @respx.mock
    async def test_aresolve_content_text_file(self):
        md_content = b"# Shared doc\n\nContent here."
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_TEXT_FILE)
        )
        respx.get(url__regex=r"/shares/u!.*/driveItem/content$").mock(
            return_value=httpx.Response(200, content=md_content)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aresolve_sharing_link_content(client, SAMPLE_SHARING_URL)
        assert item["name"] == "notes.md"
        assert content == "# Shared doc\n\nContent here."

    @respx.mock
    async def test_aresolve_content_binary_file(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_DRIVE_ITEM)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aresolve_sharing_link_content(client, SAMPLE_SHARING_URL)
        assert item["name"] == "Q4-Presentation.pptx"
        assert content is None

    @respx.mock
    async def test_alist_sharing_link_children(self):
        respx.get(url__regex=r"/shares/u!.*/root/children").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_FOLDER_CHILDREN)
        )
        async with AsyncGraphClient("tok") as client:
            items = await files.alist_sharing_link_children(client, SAMPLE_SHARING_URL)
        assert len(items) == 2
        assert items[0]["name"] == "file1.docx"


# ---------------------------------------------------------------------------
# Document extraction integration tests
# ---------------------------------------------------------------------------

SAMPLE_DOCX_ITEM = {
    "id": "file-id-docx-001",
    "name": "report.docx",
    "size": 50_000,
    "file": {"mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "lastModifiedDateTime": "2025-12-20T09:00:00Z",
    "lastModifiedBy": {"user": {"displayName": "Alice Smith", "id": "user-001"}},
    "webUrl": "https://contoso.sharepoint.com/sites/eng/report.docx",
    "parentReference": {"driveId": "drive-001", "path": "/drive/root:/Documents"},
}


def _make_sample_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Integration Test", level=1)
    doc.add_paragraph("This is test content for extraction.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocumentExtractionIntegration:
    """Integration tests for document content extraction via Graph API."""

    @respx.mock
    async def test_aget_drive_item_extracted_content_docx(self):
        docx_bytes = _make_sample_docx()
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-docx-001").mock(
            return_value=httpx.Response(200, json=SAMPLE_DOCX_ITEM)
        )
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-docx-001/content").mock(
            return_value=httpx.Response(200, content=docx_bytes)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_extracted_content(
                client, "file-id-docx-001"
            )
        assert item["name"] == "report.docx"
        assert content is not None
        assert "Integration Test" in content
        assert "This is test content" in content

    @respx.mock
    async def test_aget_drive_item_extracted_content_image_returns_none(self):
        """Image files are not extractable — returns None."""
        image_item = {
            "id": "file-id-img-001",
            "name": "photo.png",
            "size": 500_000,
            "file": {"mimeType": "image/png"},
            "lastModifiedDateTime": "2025-12-20T09:00:00Z",
            "lastModifiedBy": {"user": {"displayName": "Bob", "id": "user-002"}},
            "webUrl": "https://contoso.sharepoint.com/photo.png",
            "parentReference": {"driveId": "drive-001"},
        }
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-img-001").mock(
            return_value=httpx.Response(200, json=image_item)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_extracted_content(client, "file-id-img-001")
        assert item["name"] == "photo.png"
        assert content is None

    @respx.mock
    async def test_aget_drive_item_extracted_content_too_large(self):
        """Files exceeding 50 MB return None without downloading."""
        huge_item = {
            **SAMPLE_DOCX_ITEM,
            "id": "file-id-huge",
            "size": 60_000_000,
        }
        respx.get(f"{GRAPH_BASE_URL}/me/drive/items/file-id-huge").mock(
            return_value=httpx.Response(200, json=huge_item)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aget_drive_item_extracted_content(client, "file-id-huge")
        assert content is None

    @respx.mock
    async def test_aresolve_sharing_link_extracted_content(self):
        """Document extraction works through sharing link resolution."""
        docx_bytes = _make_sample_docx()
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_DOCX_ITEM)
        )
        respx.get(url__regex=r"/shares/u!.*/driveItem/content$").mock(
            return_value=httpx.Response(200, content=docx_bytes)
        )
        async with AsyncGraphClient("tok") as client:
            item, content = await files.aresolve_sharing_link_extracted_content(
                client, SAMPLE_SHARING_URL
            )
        assert item["name"] == "report.docx"
        assert content is not None
        assert "Integration Test" in content


# ---------------------------------------------------------------------------
# Upload sessions
# ---------------------------------------------------------------------------

_SESSION_URL = f"{GRAPH_BASE_URL}/me/drive/root:/Attachments/big.bin:/createUploadSession"


def _ranges(route):
    """Content-Range headers of every fragment PUT, in order."""
    return [call.request.headers["content-range"] for call in route.calls]


class TestUploadSessionAsync:
    """aupload_bytes_session drives a resumable upload."""

    @respx.mock
    async def test_uploads_every_fragment_in_order(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        session = respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        put = respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            side_effect=[
                httpx.Response(202, json={"nextExpectedRanges": ["4-9"]}),
                httpx.Response(202, json={"nextExpectedRanges": ["8-9"]}),
                httpx.Response(201, json=SAMPLE_UPLOADED_FILE),
            ]
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.aupload_bytes_session(
                client, "Attachments", "big.bin", b"0123456789"
            )

        assert item["id"] == SAMPLE_UPLOADED_FILE["id"]
        assert put.call_count == 3
        assert _ranges(put) == ["bytes 0-3/10", "bytes 4-7/10", "bytes 8-9/10"]
        assert json.loads(session.calls[0].request.content) == {
            "item": {"@microsoft.graph.conflictBehavior": "replace", "name": "big.bin"}
        }
        assert b"0123" == put.calls[0].request.content

    @respx.mock
    async def test_fragment_failure_cancels_session_and_reraises(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            side_effect=[
                httpx.Response(202, json={"nextExpectedRanges": ["4-9"]}),
                httpx.Response(500, json={"error": {"code": "ServiceError"}}),
            ]
        )
        cancel = respx.delete(SAMPLE_DRIVE_UPLOAD_URL).mock(return_value=httpx.Response(204))

        with pytest.raises(GraphError) as exc:
            async with AsyncGraphClient("tok") as client:
                await files.aupload_bytes_session(client, "Attachments", "big.bin", b"0123456789")

        assert exc.value.status_code == 500
        assert cancel.call_count == 1

    @respx.mock
    async def test_cancel_failure_does_not_mask_the_original_error(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            return_value=httpx.Response(500, json={"error": {"code": "ServiceError"}})
        )
        respx.delete(SAMPLE_DRIVE_UPLOAD_URL).mock(return_value=httpx.Response(404, json={}))

        with pytest.raises(GraphError) as exc:
            async with AsyncGraphClient("tok") as client:
                await files.aupload_bytes_session(client, "Attachments", "big.bin", b"0123456789")

        assert exc.value.status_code == 500

    @respx.mock
    async def test_session_without_upload_url_raises(self):
        respx.post(_SESSION_URL).mock(return_value=httpx.Response(200, json={}))
        with pytest.raises(GraphError) as exc:
            async with AsyncGraphClient("tok") as client:
                await files.aupload_bytes_session(client, "Attachments", "big.bin", b"data")

        assert exc.value.error_code == "NoUploadUrl"

    async def test_empty_payload_is_refused(self):
        async with AsyncGraphClient("tok") as client:
            with pytest.raises(ValueError, match="empty content"):
                await files.aupload_bytes_session(client, "Attachments", "big.bin", b"")

    @respx.mock
    async def test_parent_id_uses_item_path(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 16)
        session = respx.post(
            f"{GRAPH_BASE_URL}/drives/drive-001/items/parent-9:/big.bin:/createUploadSession"
        ).mock(return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION))
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            await files.aupload_bytes_session(
                client,
                "",
                "big.bin",
                b"0123456789",
                drive_id="drive-001",
                parent_id="parent-9",
                conflict_behavior="rename",
            )

        assert session.called
        assert (
            json.loads(session.calls[0].request.content)["item"][
                "@microsoft.graph.conflictBehavior"
            ]
            == "rename"
        )

    @respx.mock
    async def test_root_path_form_without_folder(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 16)
        session = respx.post(f"{GRAPH_BASE_URL}/me/drive/root:/big.bin:/createUploadSession").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            await files.aupload_bytes_session(client, "", "big.bin", b"0123456789")

        assert session.called


class TestUploadSessionSync:
    """Synchronous twin of the session upload."""

    @respx.mock
    def test_uploads_every_fragment_in_order(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        put = respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            side_effect=[
                httpx.Response(202, json={"nextExpectedRanges": ["4-9"]}),
                httpx.Response(202, json={"nextExpectedRanges": ["8-9"]}),
                httpx.Response(201, json=SAMPLE_UPLOADED_FILE),
            ]
        )
        with GraphClient("tok") as client:
            item = files.upload_bytes_session(client, "Attachments", "big.bin", b"0123456789")

        assert item["id"] == SAMPLE_UPLOADED_FILE["id"]
        assert _ranges(put) == ["bytes 0-3/10", "bytes 4-7/10", "bytes 8-9/10"]

    @respx.mock
    def test_fragment_failure_cancels_session(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            return_value=httpx.Response(503, json={"error": {"code": "ServiceUnavailable"}})
        )
        cancel = respx.delete(SAMPLE_DRIVE_UPLOAD_URL).mock(return_value=httpx.Response(204))

        with pytest.raises(GraphError):
            with GraphClient("tok") as client:
                files.upload_bytes_session(client, "Attachments", "big.bin", b"0123456789")

        assert cancel.call_count == 1


class TestUploadAny:
    """upload_any picks simple upload or a session at the 4 MB boundary."""

    @respx.mock
    async def test_at_the_limit_uses_simple_put(self, monkeypatch):
        monkeypatch.setattr(files, "MAX_SIMPLE_UPLOAD_BYTES", 5)
        simple = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/Attachments/big.bin:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            item = await files.aupload_any(client, "Attachments", "big.bin", b"01234")

        assert item["id"] == SAMPLE_UPLOADED_FILE["id"]
        assert simple.call_count == 1
        assert simple.calls[0].request.headers["Content-Type"] == "application/octet-stream"

    @respx.mock
    async def test_one_byte_over_the_limit_uses_a_session(self, monkeypatch):
        monkeypatch.setattr(files, "MAX_SIMPLE_UPLOAD_BYTES", 5)
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 4)
        session = respx.post(_SESSION_URL).mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION)
        )
        put = respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            side_effect=[
                httpx.Response(202, json={"nextExpectedRanges": ["4-5"]}),
                httpx.Response(201, json=SAMPLE_UPLOADED_FILE),
            ]
        )
        async with AsyncGraphClient("tok") as client:
            await files.aupload_any(client, "Attachments", "big.bin", b"012345")

        assert session.call_count == 1
        assert _ranges(put) == ["bytes 0-3/6", "bytes 4-5/6"]

    @respx.mock
    async def test_conflict_behavior_is_a_query_param_only_when_not_replace(self):
        route = respx.put(
            url__startswith=f"{GRAPH_BASE_URL}/me/drive/root:/Attachments/note.txt:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        async with AsyncGraphClient("tok") as client:
            await files.aupload_any(client, "Attachments", "note.txt", b"hi")
            await files.aupload_any(
                client, "Attachments", "note.txt", b"hi", conflict_behavior="rename"
            )

        assert "conflictBehavior" not in str(route.calls[0].request.url)
        assert "@microsoft.graph.conflictBehavior=rename" in str(route.calls[1].request.url)

    @respx.mock
    async def test_parent_id_simple_upload_path(self):
        route = respx.put(
            f"{GRAPH_BASE_URL}/drives/drive-001/items/parent-9:/note.txt:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        async with AsyncGraphClient("tok") as client:
            await files.aupload_any(
                client, "", "note.txt", b"hi", drive_id="drive-001", parent_id="parent-9"
            )

        assert route.called

    @respx.mock
    def test_sync_simple_upload(self):
        route = respx.put(f"{GRAPH_BASE_URL}/me/drive/root:/note.txt:/content").mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        with GraphClient("tok") as client:
            files.upload_any(client, "", "note.txt", b"hi", content_type="text/plain")

        assert route.calls[0].request.headers["Content-Type"] == "text/plain"


# ---------------------------------------------------------------------------
# Sharing-link bytes and thumbnails
# ---------------------------------------------------------------------------


class TestSharingLinkBytes:
    """resolve_sharing_link_bytes downloads any file, text or binary."""

    @respx.mock
    async def test_returns_item_and_bytes(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_DRIVE_ITEM)
        )
        respx.get(url__regex=r"/shares/u!.*/driveItem/content$").mock(
            return_value=httpx.Response(200, content=b"PPTXBYTES")
        )
        async with AsyncGraphClient("tok") as client:
            item, data = await files.aresolve_sharing_link_bytes(client, SAMPLE_SHARING_URL)

        assert item["name"] == "Q4-Presentation.pptx"
        assert data == b"PPTXBYTES"

    @respx.mock
    async def test_folder_is_refused(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_DRIVE_ITEM_FOLDER)
        )
        async with AsyncGraphClient("tok") as client:
            with pytest.raises(ValueError, match="folder"):
                await files.aresolve_sharing_link_bytes(client, SAMPLE_SHARING_URL)

    @respx.mock
    async def test_over_cap_is_refused(self):
        huge = {**SAMPLE_SHARED_DRIVE_ITEM, "size": 60_000_000}
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=huge)
        )
        async with AsyncGraphClient("tok") as client:
            with pytest.raises(ValueError, match="exceeds"):
                await files.aresolve_sharing_link_bytes(client, SAMPLE_SHARING_URL)

    @respx.mock
    def test_sync_returns_item_and_bytes(self):
        respx.get(url__regex=r"/shares/u!.*/driveItem$").mock(
            return_value=httpx.Response(200, json=SAMPLE_SHARED_TEXT_FILE)
        )
        respx.get(url__regex=r"/shares/u!.*/driveItem/content$").mock(
            return_value=httpx.Response(200, content=b"# notes")
        )
        with GraphClient("tok") as client:
            item, data = files.resolve_sharing_link_bytes(client, SAMPLE_SHARING_URL)

        assert (item["name"], data) == ("notes.md", b"# notes")


class TestSharingLinkThumbnail:
    """Thumbnails answer with bytes plus a content type, or nothing at all."""

    @respx.mock
    async def test_returns_bytes_and_type(self):
        route = respx.get(url__regex=r"/shares/u!.*/driveItem/thumbnails/0/medium/content$").mock(
            return_value=httpx.Response(
                200, content=b"THUMB", headers={"Content-Type": "image/jpeg"}
            )
        )
        async with AsyncGraphClient("tok") as client:
            result = await files.aget_sharing_link_thumbnail(client, SAMPLE_SHARING_URL)

        assert result == (b"THUMB", "image/jpeg")
        assert route.called

    @respx.mock
    async def test_404_means_no_thumbnail(self):
        respx.get(url__regex=r"/shares/u!.*/thumbnails/0/large/content$").mock(
            return_value=httpx.Response(404, json=GRAPH_ERROR_404)
        )
        async with AsyncGraphClient("tok") as client:
            assert (
                await files.aget_sharing_link_thumbnail(client, SAMPLE_SHARING_URL, size="large")
                is None
            )

    @respx.mock
    async def test_other_errors_propagate(self):
        respx.get(url__regex=r"/shares/u!.*/thumbnails/0/small/content$").mock(
            return_value=httpx.Response(403, json=GRAPH_ERROR_403)
        )
        async with AsyncGraphClient("tok") as client:
            with pytest.raises(GraphError) as exc:
                await files.aget_sharing_link_thumbnail(client, SAMPLE_SHARING_URL, size="small")

        assert exc.value.status_code == 403

    async def test_bad_size_is_refused_before_any_request(self):
        async with AsyncGraphClient("tok") as client:
            with pytest.raises(ValueError, match="small"):
                await files.aget_sharing_link_thumbnail(client, SAMPLE_SHARING_URL, size="huge")

    @respx.mock
    def test_sync_returns_bytes_and_type(self):
        respx.get(url__regex=r"/shares/u!.*/thumbnails/0/medium/content$").mock(
            return_value=httpx.Response(
                200, content=b"THUMB", headers={"Content-Type": "image/jpeg"}
            )
        )
        with GraphClient("tok") as client:
            assert files.get_sharing_link_thumbnail(client, SAMPLE_SHARING_URL) == (
                b"THUMB",
                "image/jpeg",
            )

    def test_sync_bad_size_is_refused(self):
        with GraphClient("tok") as client:
            with pytest.raises(ValueError):
                files.get_sharing_link_thumbnail(client, SAMPLE_SHARING_URL, size="tiny")


class TestUploadPathEncoding:
    """Names from callers may carry '#', '?' or spaces; the path form must survive them."""

    @respx.mock
    async def test_simple_upload_encodes_file_and_folder_names(self):
        route = respx.put(
            f"{GRAPH_BASE_URL}/me/drive/root:/Team%20Files/Q3%20%231%3F.pdf:/content"
        ).mock(return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE))
        async with AsyncGraphClient("tok") as client:
            await files.aupload_any(client, "Team Files", "Q3 #1?.pdf", b"%PDF")

        assert route.called

    @respx.mock
    async def test_session_path_encodes_file_name(self, monkeypatch):
        monkeypatch.setattr(files, "UPLOAD_CHUNK_BYTES", 16)
        session = respx.post(
            f"{GRAPH_BASE_URL}/me/drive/root:/a%23b.bin:/createUploadSession"
        ).mock(return_value=httpx.Response(200, json=SAMPLE_DRIVE_UPLOAD_SESSION))
        respx.put(SAMPLE_DRIVE_UPLOAD_URL).mock(
            return_value=httpx.Response(201, json=SAMPLE_UPLOADED_FILE)
        )
        async with AsyncGraphClient("tok") as client:
            await files.aupload_bytes_session(client, "", "a#b.bin", b"0123456789")

        assert session.called

    def test_slashes_stay_in_folders_and_leave_file_names(self):
        assert files._path_segments("/Shared Documents/Sub/", "x/y.txt") == (
            "Shared%20Documents/Sub",
            "x%2Fy.txt",
        )
