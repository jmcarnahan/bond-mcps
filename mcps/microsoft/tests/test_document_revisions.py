"""Tests for document_revisions.py — Open XML track changes manipulation."""

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from ms_graph import document_revisions


def _make_doc_with_text(*paragraphs: str) -> Document:
    """Create a Document with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


def _make_para_with_runs(*run_texts: str) -> etree._Element:
    """Create a paragraph element with multiple runs (simulates Word splitting text)."""
    doc = Document()
    para = doc.add_paragraph()
    para.clear()
    for text in run_texts:
        para.add_run(text)
    return para._element


def _make_bold_para_with_runs(runs: list[tuple[str, bool]]) -> etree._Element:
    """Create a paragraph with runs that have varying bold formatting."""
    doc = Document()
    para = doc.add_paragraph()
    para.clear()
    for text, bold in runs:
        run = para.add_run(text)
        run.bold = bold
    return para._element


class TestGetParagraphText:
    def test_single_run(self):
        para = _make_para_with_runs("Hello world")
        assert document_revisions.get_paragraph_text(para) == "Hello world"

    def test_multiple_runs(self):
        para = _make_para_with_runs("Hello ", "world", "!")
        assert document_revisions.get_paragraph_text(para) == "Hello world!"

    def test_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        assert document_revisions.get_paragraph_text(para._element) == ""

    def test_reads_text_inside_ins(self):
        para = _make_para_with_runs("before ")
        # Add an <w:ins> element with a run
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Test")
        ins.set(qn("w:date"), "2026-01-01T00:00:00Z")
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "inserted"
        run.append(t)
        ins.append(run)
        para.append(ins)

        assert document_revisions.get_paragraph_text(para) == "before inserted"

    def test_skips_text_inside_del(self):
        para = _make_para_with_runs("visible ")
        # Add a <w:del> element
        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), "1")
        del_elem.set(qn("w:author"), "Test")
        run = OxmlElement("w:r")
        dt = OxmlElement("w:delText")
        dt.text = "deleted"
        run.append(dt)
        del_elem.append(run)
        para.append(del_elem)

        assert document_revisions.get_paragraph_text(para) == "visible "


class TestFindRunsForRange:
    def test_single_run_full_match(self):
        para = _make_para_with_runs("Hello")
        result = document_revisions.find_runs_for_range(para, 0, 5)
        assert len(result) == 1
        assert result[0]["text_before"] == ""
        assert result[0]["text_match"] == "Hello"
        assert result[0]["text_after"] == ""

    def test_single_run_partial_match(self):
        para = _make_para_with_runs("Hello world")
        result = document_revisions.find_runs_for_range(para, 6, 11)
        assert len(result) == 1
        assert result[0]["text_before"] == "Hello "
        assert result[0]["text_match"] == "world"
        assert result[0]["text_after"] == ""

    def test_multi_run_span(self):
        para = _make_para_with_runs("He", "llo wor", "ld")
        # "llo wor" starts at offset 2, "ld" starts at offset 9
        # Find "lo world" = offset 3 to 11
        result = document_revisions.find_runs_for_range(para, 3, 11)
        assert len(result) == 2
        assert result[0]["text_before"] == "l"
        assert result[0]["text_match"] == "lo wor"
        assert result[0]["text_after"] == ""
        assert result[1]["text_before"] == ""
        assert result[1]["text_match"] == "ld"
        assert result[1]["text_after"] == ""

    def test_boundary_alignment(self):
        para = _make_para_with_runs("ABC", "DEF", "GHI")
        # Match "DEF" exactly (offset 3 to 6)
        result = document_revisions.find_runs_for_range(para, 3, 6)
        assert len(result) == 1
        assert result[0]["text_before"] == ""
        assert result[0]["text_match"] == "DEF"
        assert result[0]["text_after"] == ""

    def test_span_three_runs(self):
        para = _make_para_with_runs("AB", "CD", "EF")
        # Match "BCDE" = offset 1 to 5
        result = document_revisions.find_runs_for_range(para, 1, 5)
        assert len(result) == 3
        assert result[0]["text_match"] == "B"
        assert result[1]["text_match"] == "CD"
        assert result[2]["text_match"] == "E"

    def test_skips_zero_text_runs(self):
        """Runs with no w:t (e.g., containing only w:br) should be excluded."""
        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        para.add_run("Hel")
        # Insert a run with only a line break (no text)
        br_run = OxmlElement("w:r")
        br_run.append(OxmlElement("w:br"))
        para._element.append(br_run)
        para.add_run("lo world")

        result = document_revisions.find_runs_for_range(para._element, 0, 5)
        # Should only find the two text runs, not the br run
        assert len(result) == 2
        assert result[0]["text_match"] == "Hel"
        assert result[1]["text_match"] == "lo"


class TestReplaceWithRevision:
    def setup_method(self):
        document_revisions.reset_revision_counter()

    def test_simple_replace_single_run(self):
        para = _make_para_with_runs("Hello world")
        result = document_revisions.replace_with_revision(
            para, "world", "earth", "Test Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        # Should contain w:del and w:ins elements
        xml = etree.tostring(para, encoding="unicode")
        assert "w:del" in xml
        assert "w:ins" in xml
        assert "world" in xml  # in delText
        assert "earth" in xml  # in ins/r/t

    def test_replace_text_not_found(self):
        para = _make_para_with_runs("Hello world")
        result = document_revisions.replace_with_revision(
            para, "missing", "replacement", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is False

    def test_replace_spanning_two_runs(self):
        para = _make_para_with_runs("Hello ", "world!")
        result = document_revisions.replace_with_revision(
            para, "o world", "o earth", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        xml = etree.tostring(para, encoding="unicode")
        # Deleted text is split across two delText elements (one per original run)
        assert "o " in xml and "world" in xml
        assert "o earth" in xml  # inserted text
        # Verify the del element contains both parts
        del_elem = para.find(f".//{qn('w:del')}")
        del_texts = [dt.text for dt in del_elem.findall(f".//{qn('w:delText')}")]
        assert del_texts == ["o ", "world"]

    def test_replace_spanning_three_runs(self):
        para = _make_para_with_runs("AB", "CD", "EF")
        result = document_revisions.replace_with_revision(
            para, "BCDE", "XY", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        xml = etree.tostring(para, encoding="unicode")
        assert "XY" in xml

    def test_replace_at_run_start(self):
        para = _make_para_with_runs("Hello world")
        result = document_revisions.replace_with_revision(
            para, "Hello", "Hi", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        xml = etree.tostring(para, encoding="unicode")
        assert "Hello" in xml  # in delText
        assert "Hi" in xml  # in ins

    def test_replace_at_run_end(self):
        para = _make_para_with_runs("Hello world")
        result = document_revisions.replace_with_revision(
            para, "world", "earth", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True

    def test_replace_entire_paragraph_text(self):
        para = _make_para_with_runs("Replace me entirely")
        result = document_revisions.replace_with_revision(
            para, "Replace me entirely", "New text", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        xml = etree.tostring(para, encoding="unicode")
        assert "New text" in xml

    def test_replace_preserves_formatting(self):
        para = _make_bold_para_with_runs([("Hello ", False), ("bold", True), (" world", False)])
        # Replace "bold" which is in a bold run
        result = document_revisions.replace_with_revision(
            para, "bold", "BOLD", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        # The ins element should have rPr with bold
        xml = etree.tostring(para, encoding="unicode")
        ins_elem = para.find(f".//{qn('w:ins')}")
        assert ins_elem is not None
        ins_run = ins_elem.find(qn("w:r"))
        rpr = ins_run.find(qn("w:rPr"))
        assert rpr is not None
        b_elem = rpr.find(qn("w:b"))
        assert b_elem is not None

    def test_replace_produces_valid_docx(self):
        """Round-trip: modify and save as valid .docx."""
        doc = _make_doc_with_text("Hello world", "Second paragraph")
        para = doc.paragraphs[0]._element
        document_revisions.replace_with_revision(
            para, "world", "earth", "Author", "2026-01-01T00:00:00Z"
        )
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        # Should be loadable
        doc2 = Document(buf)
        assert len(doc2.paragraphs) >= 2

    def test_author_and_date_in_markup(self):
        para = _make_para_with_runs("Hello world")
        document_revisions.replace_with_revision(
            para, "world", "earth", "Bond AI", "2026-06-20T12:00:00Z"
        )
        del_elem = para.find(f".//{qn('w:del')}")
        assert del_elem.get(qn("w:author")) == "Bond AI"
        assert del_elem.get(qn("w:date")) == "2026-06-20T12:00:00Z"
        ins_elem = para.find(f".//{qn('w:ins')}")
        assert ins_elem.get(qn("w:author")) == "Bond AI"

    def test_replace_with_empty_string(self):
        """Replacing with empty string produces only w:del, no w:ins."""
        para = _make_para_with_runs("Remove this word")
        result = document_revisions.replace_with_revision(
            para, "this ", "", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True
        assert para.find(f".//{qn('w:del')}") is not None
        assert para.find(f".//{qn('w:ins')}") is None


class TestInsertWithRevision:
    def setup_method(self):
        document_revisions.reset_revision_counter()

    def test_insert_at_end(self):
        para = _make_para_with_runs("Hello")
        document_revisions.insert_with_revision(
            para, " world", "Author", "2026-01-01T00:00:00Z", position="end"
        )
        ins = para.find(f".//{qn('w:ins')}")
        assert ins is not None
        t = ins.find(f".//{qn('w:t')}")
        assert t.text == " world"

    def test_insert_at_start(self):
        para = _make_para_with_runs("world")
        document_revisions.insert_with_revision(
            para, "Hello ", "Author", "2026-01-01T00:00:00Z", position="start"
        )
        # w:ins should be the first child
        first_child = para[0]
        assert etree.QName(first_child.tag).localname == "ins"


class TestDeleteParagraphWithRevision:
    def setup_method(self):
        document_revisions.reset_revision_counter()

    def test_single_run_paragraph(self):
        para = _make_para_with_runs("Delete me")
        document_revisions.delete_paragraph_with_revision(para, "Author", "2026-01-01T00:00:00Z")
        del_elem = para.find(qn("w:del"))
        assert del_elem is not None
        del_text = del_elem.find(f".//{qn('w:delText')}")
        assert del_text.text == "Delete me"

    def test_multi_run_paragraph(self):
        para = _make_para_with_runs("Hello ", "world")
        document_revisions.delete_paragraph_with_revision(para, "Author", "2026-01-01T00:00:00Z")
        del_elem = para.find(qn("w:del"))
        assert del_elem is not None
        del_texts = del_elem.findall(f".//{qn('w:delText')}")
        assert len(del_texts) == 2
        assert del_texts[0].text == "Hello "
        assert del_texts[1].text == "world"

    def test_preserves_paragraph_element(self):
        doc = _make_doc_with_text("Delete me")
        para = doc.paragraphs[0]._element
        document_revisions.delete_paragraph_with_revision(para, "Author", "2026-01-01T00:00:00Z")
        # Paragraph still in document body
        body = doc.element.find(qn("w:body"))
        assert para in list(body)


class TestEnableTrackChanges:
    def test_adds_track_changes_to_settings(self):
        doc = Document()
        document_revisions.enable_track_changes(doc)
        settings = doc.settings.element
        tc = settings.find(qn("w:trackChanges"))
        assert tc is not None

    def test_idempotent(self):
        doc = Document()
        document_revisions.enable_track_changes(doc)
        document_revisions.enable_track_changes(doc)
        settings = doc.settings.element
        tcs = settings.findall(qn("w:trackChanges"))
        assert len(tcs) == 1
