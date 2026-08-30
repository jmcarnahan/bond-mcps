"""Tests for document_edit.py — high-level edit orchestration."""

import io
import json

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from ms_graph import document_edit


def _make_docx_bytes(*paragraphs: str) -> bytes:
    """Create a .docx in memory with the given paragraphs."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestParseEdits:
    def test_valid_replace(self):
        ops = document_edit.parse_edits(
            json.dumps([{"op": "replace", "find": "old", "replace": "new"}])
        )
        assert len(ops) == 1
        assert ops[0]["op"] == "replace"

    def test_valid_append(self):
        ops = document_edit.parse_edits(json.dumps([{"op": "append", "content": "new para"}]))
        assert ops[0]["op"] == "append"

    def test_valid_insert_after(self):
        ops = document_edit.parse_edits(
            json.dumps([{"op": "insert_after", "after": "target", "content": "new"}])
        )
        assert ops[0]["op"] == "insert_after"

    def test_valid_delete(self):
        ops = document_edit.parse_edits(json.dumps([{"op": "delete", "find": "remove me"}]))
        assert ops[0]["op"] == "delete"

    def test_valid_comment(self):
        ops = document_edit.parse_edits(
            json.dumps([{"op": "comment", "find": "text", "comment": "note"}])
        )
        assert ops[0]["op"] == "comment"

    def test_multiple_operations(self):
        ops = document_edit.parse_edits(
            json.dumps(
                [
                    {"op": "replace", "find": "a", "replace": "b"},
                    {"op": "append", "content": "new"},
                    {"op": "comment", "find": "x", "comment": "y"},
                ]
            )
        )
        assert len(ops) == 3

    def test_invalid_json(self):
        with pytest.raises(ValueError, match="Invalid JSON"):
            document_edit.parse_edits("not json{")

    def test_not_array(self):
        with pytest.raises(ValueError, match="must be a JSON array"):
            document_edit.parse_edits(json.dumps({"op": "replace"}))

    def test_missing_op(self):
        with pytest.raises(ValueError, match="missing 'op'"):
            document_edit.parse_edits(json.dumps([{"find": "x"}]))

    def test_unknown_op(self):
        with pytest.raises(ValueError, match="unknown op"):
            document_edit.parse_edits(json.dumps([{"op": "unknown"}]))

    def test_missing_required_field(self):
        with pytest.raises(ValueError, match="missing or invalid 'replace'"):
            document_edit.parse_edits(json.dumps([{"op": "replace", "find": "old"}]))


class TestApplyEdits:
    def test_replace_with_track_changes(self):
        doc_bytes = _make_docx_bytes("Hello world", "Second paragraph")
        ops = [{"op": "replace", "find": "world", "replace": "earth"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:del" in body_xml
        assert "w:ins" in body_xml
        assert "earth" in body_xml

    def test_replace_without_track_changes(self):
        doc_bytes = _make_docx_bytes("Hello world")
        ops = [{"op": "replace", "find": "world", "replace": "earth"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        assert "earth" in doc.paragraphs[0].text
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:del" not in body_xml

    def test_append_paragraph(self):
        doc_bytes = _make_docx_bytes("First")
        ops = [{"op": "append", "content": "Second paragraph"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "Second paragraph" in texts

    def test_append_with_track_changes(self):
        doc_bytes = _make_docx_bytes("First")
        ops = [{"op": "append", "content": "New content"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:ins" in body_xml
        assert "New content" in body_xml

    def test_insert_after_paragraph(self):
        doc_bytes = _make_docx_bytes("First", "Second", "Third")
        ops = [{"op": "insert_after", "after": "Second", "content": "Inserted"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text]
        second_idx = texts.index("Second")
        assert texts[second_idx + 1] == "Inserted"

    def test_delete_paragraph(self):
        doc_bytes = _make_docx_bytes("Keep", "Delete me", "Also keep")
        ops = [{"op": "delete", "find": "Delete me"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "Delete me" not in texts
        assert "Keep" in texts
        assert "Also keep" in texts

    def test_delete_with_track_changes(self):
        doc_bytes = _make_docx_bytes("Keep", "Delete me")
        ops = [{"op": "delete", "find": "Delete me"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "w:del" in body_xml
        assert "Delete me" in body_xml  # preserved in delText

    def test_comment_operation(self):
        doc_bytes = _make_docx_bytes("Hello world, please review.")
        ops = [{"op": "comment", "find": "please review", "comment": "Done!"}]

        result = document_edit.apply_edits(doc_bytes, ops)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "commentRangeStart" in body_xml
        assert "commentRangeEnd" in body_xml

    def test_multiple_edits_in_sequence(self):
        doc_bytes = _make_docx_bytes("Hello world", "Second paragraph")
        ops = [
            {"op": "replace", "find": "world", "replace": "earth"},
            {"op": "append", "content": "Third paragraph"},
        ]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "earth" in body_xml
        assert "Third paragraph" in body_xml

    def test_edit_error_on_text_not_found(self):
        doc_bytes = _make_docx_bytes("Hello world")
        ops = [{"op": "replace", "find": "missing text", "replace": "new"}]

        with pytest.raises(document_edit.EditError, match="Text not found"):
            document_edit.apply_edits(doc_bytes, ops)

    def test_output_is_valid_docx(self):
        doc_bytes = _make_docx_bytes("Test document")
        ops = [{"op": "replace", "find": "Test", "replace": "Modified"}]

        result = document_edit.apply_edits(doc_bytes, ops)
        # Should be loadable as a valid docx
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_custom_author(self):
        doc_bytes = _make_docx_bytes("Hello world")
        ops = [{"op": "replace", "find": "world", "replace": "earth"}]

        result = document_edit.apply_edits(
            doc_bytes, ops, track_changes=True, author="Custom Author"
        )

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "Custom Author" in body_xml


class TestApplyEditsRoundTrip:
    """Validate that edited documents can be re-read and contain correct content."""

    def test_replace_visible_text_after_accepting_changes(self):
        """After accepting all track changes, the replacement text should be visible."""
        doc_bytes = _make_docx_bytes("The quick brown fox jumps over the lazy dog.")
        ops = [{"op": "replace", "find": "brown fox", "replace": "red cat"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        # Load and check: the ins text should contain 'red cat'
        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        # The insertion text is what would remain after accepting
        assert "red cat" in body_xml
        # The deletion text preserves the original
        assert "brown fox" in body_xml

    def test_multiple_replaces_all_applied(self):
        doc_bytes = _make_docx_bytes("Alice went home.", "Bob stayed out.")
        ops = [
            {"op": "replace", "find": "Alice", "replace": "Carol"},
            {"op": "replace", "find": "Bob", "replace": "Dave"},
        ]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        assert "Carol" in doc.paragraphs[0].text
        assert "Dave" in doc.paragraphs[1].text
        assert "Alice" not in doc.paragraphs[0].text
        assert "Bob" not in doc.paragraphs[1].text

    def test_insert_after_preserves_document_order(self):
        doc_bytes = _make_docx_bytes("Chapter 1", "Chapter 2", "Chapter 3")
        ops = [
            {"op": "insert_after", "after": "Chapter 1", "content": "Section 1.1"},
            {"op": "insert_after", "after": "Chapter 2", "content": "Section 2.1"},
        ]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert texts == ["Chapter 1", "Section 1.1", "Chapter 2", "Section 2.1", "Chapter 3"]

    def test_delete_then_append_workflow(self):
        """Simulate removing a paragraph and adding a replacement at the end."""
        doc_bytes = _make_docx_bytes("Intro", "Old content", "Conclusion")
        ops = [
            {"op": "delete", "find": "Old content"},
            {"op": "append", "content": "New content"},
        ]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "Old content" not in texts
        assert "New content" in texts
        assert "Intro" in texts
        assert "Conclusion" in texts

    def test_comment_and_replace_together(self):
        """Comment on one phrase, replace another — both in the same document."""
        doc_bytes = _make_docx_bytes("Please review the budget for Q3 projections.")
        ops = [
            {"op": "comment", "find": "budget", "comment": "Needs updated figures"},
            {"op": "replace", "find": "Q3", "replace": "Q4"},
        ]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        # Comment markers present
        assert "commentRangeStart" in body_xml
        assert "commentRangeEnd" in body_xml
        # Revision markup present
        assert "w:del" in body_xml
        assert "Q4" in body_xml

    def test_direct_replace_preserves_bold_formatting(self):
        """Non-tracked replace should keep formatting on unaffected text."""
        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        run1 = para.add_run("Keep bold ")
        run1.bold = True
        run2 = para.add_run("replace_me")
        run2.bold = False
        run3 = para.add_run(" keep italic")
        run3.italic = True
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        ops = [{"op": "replace", "find": "replace_me", "replace": "REPLACED"}]
        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc2 = Document(io.BytesIO(result))
        para2 = doc2.paragraphs[0]
        # First run should still be bold
        assert para2.runs[0].bold is True
        assert "Keep bold " in para2.runs[0].text
        # Last run should still be italic
        assert para2.runs[2].italic is True
        assert "keep italic" in para2.runs[2].text
        # Middle run has the replacement
        assert "REPLACED" in para2.runs[1].text

    def test_error_reports_correct_operation_index(self):
        """When operation #2 fails, the error should say 'Edit #2', not 'Edit #0'."""
        doc_bytes = _make_docx_bytes("Hello world")
        ops = [
            {"op": "append", "content": "new para"},
            {"op": "append", "content": "another"},
            {"op": "replace", "find": "nonexistent", "replace": "x"},
        ]

        with pytest.raises(document_edit.EditError) as exc_info:
            document_edit.apply_edits(doc_bytes, ops)

        assert exc_info.value.op_index == 2
        assert "Edit #2" in str(exc_info.value)

    def test_direct_replace_handles_multi_t_runs(self):
        """Runs with multiple <w:t> elements should be replaced correctly."""
        from docx.oxml import OxmlElement

        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        # Build a run with TWO w:t elements: 'Hel' + 'lo'
        r = OxmlElement("w:r")
        t1 = OxmlElement("w:t")
        t1.text = "Hel"
        r.append(t1)
        t2 = OxmlElement("w:t")
        t2.text = "lo"
        r.append(t2)
        para._element.append(r)
        para.add_run(" world")

        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        ops = [{"op": "replace", "find": "Hello", "replace": "Hi"}]
        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc2 = Document(io.BytesIO(result))
        assert doc2.paragraphs[0].text == "Hi world"

    def test_replace_preserves_line_breaks_between_runs(self):
        """Line breaks (w:br) between text runs must not be destroyed by replace."""
        from docx.oxml import OxmlElement

        doc = Document()
        para = doc.add_paragraph()
        para.clear()
        para.add_run("Hel")
        # A run containing only a line break
        br_run = OxmlElement("w:r")
        br_run.append(OxmlElement("w:br"))
        para._element.append(br_run)
        para.add_run("lo world")

        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        ops = [{"op": "replace", "find": "Hello", "replace": "Hi"}]
        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc2 = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc2.paragraphs[0]._element, encoding="unicode")
        assert "w:br" in body_xml, "Line break run should survive replacement"


class TestApplyEditsEdgeCases:
    def test_empty_document(self):
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()
        ops = [{"op": "append", "content": "First content"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc2 = Document(io.BytesIO(result))
        texts = [p.text for p in doc2.paragraphs if p.text]
        assert "First content" in texts

    def test_text_appears_multiple_times_replaces_first(self):
        doc_bytes = _make_docx_bytes("Hello world", "Another world")
        ops = [{"op": "replace", "find": "world", "replace": "earth"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=False)

        doc = Document(io.BytesIO(result))
        # First occurrence replaced, second preserved
        assert "earth" in doc.paragraphs[0].text
        assert "world" in doc.paragraphs[1].text

    def test_unicode_text(self):
        doc_bytes = _make_docx_bytes("Héllo wörld café")
        ops = [{"op": "replace", "find": "wörld", "replace": "münch"}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body_xml = etree.tostring(doc.element.find(qn("w:body")), encoding="unicode")
        assert "münch" in body_xml

    def test_replace_with_empty_string(self):
        doc_bytes = _make_docx_bytes("Remove this word please")
        ops = [{"op": "replace", "find": "this ", "replace": ""}]

        result = document_edit.apply_edits(doc_bytes, ops, track_changes=True)

        doc = Document(io.BytesIO(result))
        body = doc.element.find(qn("w:body"))
        # Should have a w:del with the deleted text
        del_elem = body.find(f".//{qn('w:del')}")
        assert del_elem is not None
        del_text = del_elem.find(f".//{qn('w:delText')}")
        assert del_text is not None and del_text.text == "this "
        # No w:ins since replacement is empty
        assert body.find(f".//{qn('w:ins')}") is None
