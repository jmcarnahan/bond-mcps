"""Tests for document_comments.py — Open XML comment manipulation."""

import io

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from lxml import etree
from ms_graph import document_comments, document_revisions


def _make_doc_with_text(*paragraphs: str) -> Document:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


def _make_para_with_runs(doc: Document, *run_texts: str):
    """Add a paragraph with multiple runs to a document."""
    para = doc.add_paragraph()
    para.clear()
    for text in run_texts:
        para.add_run(text)
    return para


class TestGetOrCreateCommentsPart:
    def test_creates_part_when_missing(self):
        doc = Document()
        comments_root = document_comments.get_or_create_comments_part(doc)
        assert comments_root is not None
        assert etree.QName(comments_root.tag).localname == "comments"

    def test_returns_existing_part(self):
        doc = Document()
        root1 = document_comments.get_or_create_comments_part(doc)
        # Add a comment to make it non-empty
        comment = etree.SubElement(root1, qn("w:comment"))
        comment.set(qn("w:id"), "0")
        document_comments._save_comments_part(doc, root1)

        root2 = document_comments.get_or_create_comments_part(doc)
        # Should find the existing comment
        assert root2.find(qn("w:comment")) is not None

    def test_relationship_registered(self):
        doc = Document()
        document_comments.get_or_create_comments_part(doc)
        # Check that a COMMENTS relationship exists
        found = False
        for rel in doc.part.rels.values():
            if rel.reltype == RT.COMMENTS:
                found = True
                break
        assert found


class TestAddComment:
    def setup_method(self):
        document_revisions.reset_revision_counter()

    def test_simple_comment_on_text(self):
        doc = _make_doc_with_text("Hello world, this is a test.")
        para = doc.paragraphs[0]._element
        result = document_comments.add_comment(
            doc, para, "world", "Nice word!", "Reviewer", "2026-01-01T00:00:00Z"
        )
        assert result is True

        # Check markers in paragraph XML
        xml = etree.tostring(para, encoding="unicode")
        assert "commentRangeStart" in xml
        assert "commentRangeEnd" in xml
        assert "commentReference" in xml

    def test_comment_text_not_found(self):
        doc = _make_doc_with_text("Hello world")
        para = doc.paragraphs[0]._element
        result = document_comments.add_comment(doc, para, "missing", "Note", "Author")
        assert result is False

    def test_comment_spanning_runs(self):
        doc = Document()
        para = _make_para_with_runs(doc, "Hello ", "world!")
        result = document_comments.add_comment(
            doc, para._element, "o world", "Spans runs", "Author", "2026-01-01T00:00:00Z"
        )
        assert result is True

    def test_multiple_comments_sequential_ids(self):
        doc = _make_doc_with_text("Hello world, goodbye world.")
        para = doc.paragraphs[0]._element

        document_comments.add_comment(
            doc, para, "Hello", "First comment", "Author", "2026-01-01T00:00:00Z"
        )
        document_comments.add_comment(
            doc, para, "goodbye", "Second comment", "Author", "2026-01-01T00:00:00Z"
        )

        # Both range starts should have different IDs
        range_starts = para.findall(f".//{qn('w:commentRangeStart')}")
        ids = [rs.get(qn("w:id")) for rs in range_starts]
        assert len(ids) == 2
        assert ids[0] != ids[1]

    def test_comment_content_in_comments_xml(self):
        doc = _make_doc_with_text("Hello world")
        para = doc.paragraphs[0]._element
        document_comments.add_comment(
            doc, para, "world", "This is important", "Bond AI", "2026-01-01T00:00:00Z"
        )

        # Read comments part back
        comments_root = document_comments.get_or_create_comments_part(doc)
        comment = comments_root.find(qn("w:comment"))
        assert comment is not None
        assert comment.get(qn("w:author")) == "Bond AI"
        # Comment body text
        t_elem = comment.find(f".//{qn('w:t')}")
        assert t_elem.text == "This is important"

    def test_comment_reference_run_added(self):
        doc = _make_doc_with_text("Hello world")
        para = doc.paragraphs[0]._element
        document_comments.add_comment(doc, para, "world", "Note", "Author", "2026-01-01T00:00:00Z")
        # commentReference should exist
        ref = para.find(f".//{qn('w:commentReference')}")
        assert ref is not None

    def test_round_trip_save_load(self):
        """Document with comment should save and reload without error."""
        doc = _make_doc_with_text("Hello world", "Second paragraph")
        para = doc.paragraphs[0]._element
        document_comments.add_comment(
            doc, para, "world", "A note", "Author", "2026-01-01T00:00:00Z"
        )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)
        assert len(doc2.paragraphs) >= 2

    def test_comment_markers_in_correct_position(self):
        doc = _make_doc_with_text("ABCDEF")
        para = doc.paragraphs[0]._element
        document_comments.add_comment(doc, para, "CD", "Middle", "Author", "2026-01-01T00:00:00Z")

        # commentRangeStart should appear before the run containing "CD"
        children = list(para)
        tag_names = [etree.QName(c.tag).localname for c in children]
        start_idx = tag_names.index("commentRangeStart")
        end_idx = tag_names.index("commentRangeEnd")
        assert start_idx < end_idx

    def test_comment_full_round_trip_with_content_verification(self):
        """Save, reload, and verify comment content is intact in the package."""
        doc = _make_doc_with_text("Review this section carefully.")
        para = doc.paragraphs[0]._element
        document_comments.add_comment(
            doc, para, "this section", "Needs more detail", "Bond AI", "2026-06-20T10:00:00Z"
        )

        # Save and reload
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        # Verify comments part exists and has our comment
        from docx.opc.constants import RELATIONSHIP_TYPE as RT2

        found_comment = False
        for rel in doc2.part.rels.values():
            if rel.reltype == RT2.COMMENTS:
                comments_xml = etree.fromstring(rel.target_part.blob)
                for comment in comments_xml.findall(qn("w:comment")):
                    if comment.get(qn("w:author")) == "Bond AI":
                        t_elem = comment.find(f".//{qn('w:t')}")
                        if t_elem is not None and t_elem.text == "Needs more detail":
                            found_comment = True
                break
        assert found_comment, "Comment 'Needs more detail' by 'Bond AI' not found in saved doc"

        # Verify body still has the markers
        body_xml = etree.tostring(doc2.element.find(qn("w:body")), encoding="unicode")
        assert "commentRangeStart" in body_xml
        assert "commentRangeEnd" in body_xml

    def test_multiple_comments_on_different_paragraphs(self):
        """Comments on separate paragraphs should all persist."""
        doc = _make_doc_with_text("First paragraph.", "Second paragraph.", "Third paragraph.")

        document_comments.add_comment(
            doc, doc.paragraphs[0]._element, "First", "Comment A", "Author", "2026-01-01T00:00:00Z"
        )
        document_comments.add_comment(
            doc, doc.paragraphs[2]._element, "Third", "Comment B", "Author", "2026-01-02T00:00:00Z"
        )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        # Both paragraphs should have comment markers
        para0_xml = etree.tostring(doc2.paragraphs[0]._element, encoding="unicode")
        para2_xml = etree.tostring(doc2.paragraphs[2]._element, encoding="unicode")
        assert "commentRangeStart" in para0_xml
        assert "commentRangeStart" in para2_xml
        # Middle paragraph should NOT have comment markers
        para1_xml = etree.tostring(doc2.paragraphs[1]._element, encoding="unicode")
        assert "commentRangeStart" not in para1_xml
