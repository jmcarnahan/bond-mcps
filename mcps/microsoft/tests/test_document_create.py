"""Tests for Word document generation from markdown."""

import io

import pytest
from docx import Document
from ms_graph.document_create import (
    _apply_inline_formatting,
    csv_to_xlsx,
    markdown_to_docx,
)


class TestMarkdownToDocx:
    def _doc(self, md: str) -> Document:
        data = markdown_to_docx(md)
        return Document(io.BytesIO(data))

    def test_heading_level_1(self):
        doc = self._doc("# Hello World")
        para = doc.paragraphs[0]
        assert para.style.name == "Heading 1"
        assert para.text == "Hello World"

    def test_heading_level_2(self):
        doc = self._doc("## Section")
        para = doc.paragraphs[0]
        assert para.style.name == "Heading 2"

    def test_heading_level_3(self):
        doc = self._doc("### Subsection")
        para = doc.paragraphs[0]
        assert para.style.name == "Heading 3"

    def test_heading_level_4(self):
        doc = self._doc("#### Deep")
        para = doc.paragraphs[0]
        assert para.style.name == "Heading 4"

    def test_paragraph(self):
        doc = self._doc("Just a plain paragraph.")
        para = doc.paragraphs[0]
        assert para.style.name == "Normal"
        assert para.text == "Just a plain paragraph."

    def test_bold_formatting(self):
        doc = self._doc("This is **bold** text.")
        para = doc.paragraphs[0]
        runs = para.runs
        assert runs[0].text == "This is "
        assert runs[0].bold is not True
        assert runs[1].text == "bold"
        assert runs[1].bold is True
        assert runs[2].text == " text."

    def test_italic_star_formatting(self):
        doc = self._doc("This is *italic* text.")
        para = doc.paragraphs[0]
        runs = para.runs
        assert runs[1].text == "italic"
        assert runs[1].italic is True

    def test_underscore_treated_as_literal(self):
        doc = self._doc("This is _italic_ text.")
        para = doc.paragraphs[0]
        assert para.text == "This is _italic_ text."
        assert all(r.italic is not True for r in para.runs)

    def test_bold_italic_formatting(self):
        doc = self._doc("This is ***important*** text.")
        para = doc.paragraphs[0]
        runs = para.runs
        assert runs[1].text == "important"
        assert runs[1].bold is True
        assert runs[1].italic is True

    def test_bullet_list(self):
        md = "- First item\n- Second item\n- Third item"
        doc = self._doc(md)
        assert doc.paragraphs[0].style.name == "List Bullet"
        assert doc.paragraphs[0].text == "First item"
        assert doc.paragraphs[1].text == "Second item"
        assert doc.paragraphs[2].text == "Third item"

    def test_bullet_list_asterisk(self):
        md = "* Item one\n* Item two"
        doc = self._doc(md)
        assert doc.paragraphs[0].style.name == "List Bullet"
        assert doc.paragraphs[0].text == "Item one"

    def test_numbered_list(self):
        md = "1. First\n2. Second\n3. Third"
        doc = self._doc(md)
        assert doc.paragraphs[0].style.name == "List Number"
        assert doc.paragraphs[0].text == "First"
        assert doc.paragraphs[1].text == "Second"

    def test_table_basic(self):
        md = "| Name | Age |\n|------|-----|\n| Alice | 30 |\n| Bob | 25 |"
        doc = self._doc(md)
        tables = doc.tables
        assert len(tables) == 1
        table = tables[0]
        assert len(table.rows) == 3
        assert table.cell(0, 0).text == "Name"
        assert table.cell(0, 1).text == "Age"
        assert table.cell(1, 0).text == "Alice"
        assert table.cell(1, 1).text == "30"
        assert table.cell(2, 0).text == "Bob"
        assert table.cell(2, 1).text == "25"

    def test_table_header_bold(self):
        md = "| Col1 | Col2 |\n|------|------|\n| a | b |"
        doc = self._doc(md)
        table = doc.tables[0]
        header_para = table.cell(0, 0).paragraphs[0]
        bold_runs = [r for r in header_para.runs if r.text.strip()]
        assert bold_runs[0].bold is True

    def test_mixed_document(self):
        md = """# Contract Review

## Summary

This document reviews the **ScreenMeet** contract.

## Key Points

- Termination: 30 days notice
- Payment: Net 60
- Liability cap: $500,000

## Comparison

| Clause | Status |
|--------|--------|
| Indemnity | Acceptable |
| IP Rights | Needs review |

Final notes paragraph."""
        doc = self._doc(md)
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[0].text == "Contract Review"
        assert len(doc.tables) == 1

    def test_empty_lines_skipped(self):
        md = "First paragraph.\n\n\nSecond paragraph."
        doc = self._doc(md)
        texts = [p.text for p in doc.paragraphs if p.text]
        assert texts == ["First paragraph.", "Second paragraph."]

    def test_output_is_valid_docx(self):
        data = markdown_to_docx("# Test\n\nHello world.")
        assert data[:2] == b"PK"  # ZIP magic bytes

    def test_line_limit_exceeded(self):
        md = "\n".join(["line"] * 6000)
        with pytest.raises(ValueError, match="line limit"):
            markdown_to_docx(md)

    def test_empty_content(self):
        data = markdown_to_docx("")
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) == 0 or all(not p.text for p in doc.paragraphs)

    def test_table_does_not_mutate_normal_style(self):
        from docx.shared import Pt

        md = "Before.\n\n| A |\n|---|\n| b |\n\nAfter."
        doc = self._doc(md)
        normal_size = doc.styles["Normal"].font.size
        assert normal_size == Pt(11)

    def test_underscore_in_word_not_italic(self):
        doc = self._doc("Use file_name_here in your code.")
        para = doc.paragraphs[0]
        full_text = para.text
        assert full_text == "Use file_name_here in your code."
        assert all(r.italic is not True for r in para.runs)

    def test_lone_asterisk_preserved(self):
        doc = self._doc("Calculate 5 * 3 = 15")
        para = doc.paragraphs[0]
        assert para.text == "Calculate 5 * 3 = 15"

    def test_multiple_lone_asterisks_preserved(self):
        doc = self._doc("a * b * c")
        para = doc.paragraphs[0]
        assert para.text == "a * b * c"

    def test_dunder_preserved(self):
        doc = self._doc("The __init__.py file is special.")
        para = doc.paragraphs[0]
        assert para.text == "The __init__.py file is special."


class TestInlineFormatting:
    def _runs(self, text: str):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        _apply_inline_formatting(para, text)
        return para.runs

    def test_plain_text(self):
        runs = self._runs("Hello world")
        assert len(runs) == 1
        assert runs[0].text == "Hello world"
        assert runs[0].bold is not True

    def test_bold(self):
        runs = self._runs("**bold**")
        assert runs[0].text == "bold"
        assert runs[0].bold is True

    def test_italic_star(self):
        runs = self._runs("*italic*")
        assert runs[0].text == "italic"
        assert runs[0].italic is True

    def test_underscore_literal(self):
        runs = self._runs("_emphasis_")
        assert runs[0].text == "_emphasis_"
        assert runs[0].italic is not True

    def test_mixed_in_one_string(self):
        runs = self._runs("Normal **bold** and *italic* end")
        assert runs[0].text == "Normal "
        assert runs[1].text == "bold"
        assert runs[1].bold is True
        assert runs[2].text == " and "
        assert runs[3].text == "italic"
        assert runs[3].italic is True
        assert runs[4].text == " end"


class TestCsvToXlsx:
    def _wb(self, content: str):
        from openpyxl import load_workbook

        return load_workbook(io.BytesIO(csv_to_xlsx(content)))

    def test_empty_content_produces_valid_workbook(self):
        wb = self._wb("")
        assert wb.sheetnames == ["Sheet1"]
        # A blank sheet has no data rows.
        assert wb.active.max_row == 1
        assert wb.active["A1"].value is None

    def test_csv_rows_and_columns(self):
        wb = self._wb("Name,Qty\nWidget,10\nGadget,25")
        ws = wb.active
        assert ws["A1"].value == "Name"
        assert ws["B1"].value == "Qty"
        assert ws["A2"].value == "Widget"
        assert ws["B2"].value == 10  # coerced to int

    def test_numeric_coercion(self):
        wb = self._wb("int,float,text\n42,3.14,hello")
        ws = wb.active
        assert ws["A2"].value == 42
        assert isinstance(ws["A2"].value, int)
        assert ws["B2"].value == 3.14
        assert isinstance(ws["B2"].value, float)
        assert ws["C2"].value == "hello"

    def test_leading_zeros_preserved_as_text(self):
        """Zip codes / IDs / phone numbers must not be mangled into ints."""
        wb = self._wb("zip,id\n02134,007")
        ws = wb.active
        assert ws["A2"].value == "02134"
        assert ws["B2"].value == "007"

    def test_inf_and_nan_stay_text(self):
        """float('inf')/'nan' round-trip to empty cells in Excel — keep as text."""
        wb = self._wb("a,b,c\ninf,nan,1e999")
        ws = wb.active
        assert ws["A2"].value == "inf"
        assert ws["B2"].value == "nan"
        assert ws["C2"].value == "1e999"

    def test_python_numeric_quirks_stay_text(self):
        """Underscores and leading '+' are valid to int()/float() but not real numbers."""
        wb = self._wb("a,b\n1_000,+5")
        ws = wb.active
        assert ws["A2"].value == "1_000"
        assert ws["B2"].value == "+5"

    def test_plain_zero_is_numeric(self):
        wb = self._wb("a\n0")
        assert wb.active["A2"].value == 0

    def test_negative_and_decimal(self):
        wb = self._wb("a,b\n-42,-3.14")
        ws = wb.active
        assert ws["A2"].value == -42
        assert ws["B2"].value == -3.14

    def test_output_is_valid_xlsx(self):
        data = csv_to_xlsx("a,b\n1,2")
        assert data[:2] == b"PK"  # ZIP magic bytes

    def test_empty_middle_cell(self):
        wb = self._wb("a,b,c\n1,,3")
        ws = wb.active
        assert ws["A2"].value == 1
        assert ws["B2"].value is None
        assert ws["C2"].value == 3
