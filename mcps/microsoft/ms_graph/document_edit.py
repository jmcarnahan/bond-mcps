"""Word document editing orchestration.

Coordinates applying edit operations to a .docx file. Operations are
expressed as a JSON array of operation dicts.
"""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone

from docx import Document

from . import document_comments, document_revisions

_REQUIRED_FIELDS: dict[str, list[str]] = {
    "replace": ["find", "replace"],
    "append": ["content"],
    "insert_after": ["after", "content"],
    "delete": ["find"],
    "comment": ["find", "comment"],
}


class EditError(Exception):
    """Raised when an edit operation fails."""

    def __init__(self, op_index: int, op_type: str, message: str):
        self.op_index = op_index
        self.op_type = op_type
        super().__init__(f"Edit #{op_index} ({op_type}): {message}")


def parse_edits(edits_json: str) -> list[dict[str, str]]:
    """Parse and validate the edits JSON string."""
    try:
        data = json.loads(edits_json)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {e}")

    if not isinstance(data, list):
        raise ValueError("Edits must be a JSON array")

    for i, item in enumerate(data):
        if not isinstance(item, dict):
            raise ValueError(f"Edit #{i}: must be an object")
        op = item.get("op")
        if not op:
            raise ValueError(f"Edit #{i}: missing 'op' field")
        if op not in _REQUIRED_FIELDS:
            raise ValueError(
                f"Edit #{i}: unknown op '{op}'. " f"Valid: {', '.join(_REQUIRED_FIELDS.keys())}"
            )
        for field in _REQUIRED_FIELDS[op]:
            if field not in item or not isinstance(item[field], str):
                raise ValueError(f"Edit #{i} ({op}): missing or invalid '{field}'")

    return data


def apply_edits(
    doc_bytes: bytes,
    operations: list[dict[str, str]],
    track_changes: bool = True,
    author: str = "Bond AI",
) -> bytes:
    """Apply edit operations to a .docx document.

    Returns modified .docx as bytes.
    Raises EditError if an operation fails.
    """
    doc = Document(io.BytesIO(doc_bytes))
    document_revisions.reset_revision_counter()

    if track_changes:
        document_revisions.enable_track_changes(doc)

    date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    for i, op in enumerate(operations):
        op_type = op["op"]
        try:
            if op_type == "replace":
                _apply_replace(doc, op, i, track_changes, author, date)
            elif op_type == "append":
                _apply_append(doc, op, i, track_changes, author, date)
            elif op_type == "insert_after":
                _apply_insert_after(doc, op, i, track_changes, author, date)
            elif op_type == "delete":
                _apply_delete(doc, op, i, track_changes, author, date)
            elif op_type == "comment":
                _apply_comment(doc, op, i, author, date)
        except EditError:
            raise
        except Exception as e:
            raise EditError(i, op_type, str(e)) from e

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _apply_replace(
    doc: Document,
    op: dict[str, str],
    op_index: int,
    track_changes: bool,
    author: str,
    date: str,
) -> None:
    target = op["find"]
    replacement = op["replace"]

    for para in doc.paragraphs:
        if track_changes:
            if document_revisions.replace_with_revision(
                para._element, target, replacement, author, date
            ):
                return
        else:
            full_text = document_revisions.get_paragraph_text(para._element)
            if target in full_text:
                _direct_replace(para, target, replacement)
                return

    raise EditError(op_index, "replace", f"Text not found: '{_truncate(target)}'")


def _apply_append(
    doc: Document,
    op: dict[str, str],
    op_index: int,
    track_changes: bool,
    author: str,
    date: str,
) -> None:
    content = op["content"]
    para = doc.add_paragraph()
    if track_changes:
        document_revisions.insert_with_revision(
            para._element, content, author, date, position="end"
        )
    else:
        para.add_run(content)


def _apply_insert_after(
    doc: Document,
    op: dict[str, str],
    op_index: int,
    track_changes: bool,
    author: str,
    date: str,
) -> None:
    after_text = op["after"]
    content = op["content"]

    for para in doc.paragraphs:
        para_text = document_revisions.get_paragraph_text(para._element)
        if after_text in para_text:
            new_para = _insert_paragraph_after(para)
            if track_changes:
                document_revisions.insert_with_revision(
                    new_para._element, content, author, date, position="end"
                )
            else:
                new_para.add_run(content)
            return

    raise EditError(op_index, "insert_after", f"Text not found: '{_truncate(after_text)}'")


def _apply_delete(
    doc: Document,
    op: dict[str, str],
    op_index: int,
    track_changes: bool,
    author: str,
    date: str,
) -> None:
    target = op["find"]

    for para in doc.paragraphs:
        para_text = document_revisions.get_paragraph_text(para._element)
        if target in para_text:
            if track_changes:
                document_revisions.delete_paragraph_with_revision(para._element, author, date)
            else:
                body = para._element.getparent()
                body.remove(para._element)
            return

    raise EditError(op_index, "delete", f"Text not found: '{_truncate(target)}'")


def _apply_comment(
    doc: Document,
    op: dict[str, str],
    op_index: int,
    author: str,
    date: str,
) -> None:
    target = op["find"]
    comment_text = op["comment"]

    for para in doc.paragraphs:
        if document_comments.add_comment(doc, para._element, target, comment_text, author, date):
            return

    raise EditError(op_index, "comment", f"Text not found: '{_truncate(target)}'")


def _direct_replace(para, target: str, replacement: str) -> None:
    """Direct (non-tracked) text replacement preserving run formatting.

    Uses the same remove-and-rebuild approach as the tracked-changes path:
    removes affected run elements and inserts fresh ones built from scratch.
    This avoids in-place mutation of existing elements whose internal structure
    (multiple w:t elements, nested formatting) is unpredictable.
    """
    full_text = document_revisions.get_paragraph_text(para._element)
    idx = full_text.find(target)
    if idx == -1:
        return

    affected = document_revisions.find_runs_for_range(para._element, idx, idx + len(target))
    if not affected:
        return

    first_run = affected[0]["element"]
    first_parent = affected[0]["parent"]
    insert_pos = list(first_parent).index(first_run)

    new_elements = []
    if affected[0]["text_before"]:
        new_elements.append(document_revisions._make_run(affected[0]["text_before"], first_run))
    if replacement:
        new_elements.append(document_revisions._make_run(replacement, first_run))
    if affected[-1]["text_after"]:
        new_elements.append(
            document_revisions._make_run(affected[-1]["text_after"], affected[-1]["element"])
        )

    for entry in affected:
        entry["parent"].remove(entry["element"])

    for i, elem in enumerate(new_elements):
        first_parent.insert(insert_pos + i, elem)


def _insert_paragraph_after(para):
    """Insert a new paragraph element after the given one and return it."""
    from docx.text.paragraph import Paragraph

    new_p = _make_paragraph_element()
    para._element.addnext(new_p)
    return Paragraph(new_p, para._element.getparent())


def _make_paragraph_element():
    """Create a bare <w:p> element."""
    from docx.oxml import OxmlElement

    return OxmlElement("w:p")


def _truncate(text: str, max_len: int = 50) -> str:
    if len(text) <= max_len:
        return text
    return text[:max_len] + "..."
