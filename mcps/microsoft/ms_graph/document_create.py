"""
Word document (.docx) generation from markdown text.

Converts a subset of markdown to .docx using python-docx. All operations use
in-memory BytesIO streams — no temp files are created. Libraries are imported
lazily to keep server startup fast.

Supported markdown elements:
  - Headings: # H1 through #### H4
  - Bold: **text**
  - Italic: *text*
  - Bold+italic: ***text***
  - Bullet lists: - item or * item
  - Numbered lists: 1. item
  - Tables: pipe-delimited with header separator row
  - Paragraphs: everything else

Note: underscore italic (_text_) is intentionally NOT supported because it
causes false positives with identifiers (file_name, __init__) and technical
content. Use *asterisk italic* instead.
"""

import io
import re
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document
    from docx.text.paragraph import Paragraph

_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_BULLET_RE = re.compile(r"^[\-\*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.+)$")
_TABLE_SEP_RE = re.compile(r"^\|?\s*[-:]+[-| :]*$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")

_INLINE_RE = re.compile(
    r"\*\*\*(\S.*?\S|\S)\*\*\*" r"|\*\*(\S.*?\S|\S)\*\*" r"|\*(\S.*?\S|\S)\*" r"|([^*]+|\*)"
)

MAX_LINES = 5000


def markdown_to_docx(markdown: str) -> bytes:
    """Convert markdown text to .docx bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    lines = markdown.split("\n")
    if len(lines) > MAX_LINES:
        raise ValueError(
            f"Content has {len(lines):,} lines, exceeding the {MAX_LINES:,} line limit."
        )

    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            para = doc.add_heading(text, level=level)
            _apply_inline_formatting(para, text, clear_existing=True)
            i += 1
            continue

        if _TABLE_ROW_RE.match(line):
            i = _consume_table(doc, lines, i)
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            i = _consume_list(doc, lines, i, ordered=False)
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            i = _consume_list(doc, lines, i, ordered=True)
            continue

        para = doc.add_paragraph()
        _apply_inline_formatting(para, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def csv_to_xlsx(content: str) -> bytes:
    """Convert CSV text to a single-sheet .xlsx workbook.

    Empty content produces a valid empty workbook (one blank sheet), which is
    the common case for seeding a file that edit_excel_workbook will fill in.
    Numeric-looking cells are written as numbers so formulas and formatting
    behave; everything else stays text.
    """
    import csv

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    text = content.strip()
    if text:
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            ws.append([_coerce_cell(cell) for cell in row])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


_INT_RE = re.compile(r"^-?[0-9]+$")
_FLOAT_RE = re.compile(r"^-?[0-9]+\.[0-9]+$")


def _coerce_cell(value: str):
    """Coerce a CSV string cell to int/float when it is a plain number, else str.

    Deliberately strict — much stricter than int()/float() — to avoid silently
    corrupting data:
      - Leading zeros are preserved as text ("007" stays "007", not 7) so zip
        codes, IDs, and phone numbers survive.
      - inf/nan and scientific notation ("1e9") stay text; float("inf") would
        round-trip to an empty cell in Excel.
      - Python numeric quirks ("1_000", "+5") stay text.
    Whitespace-only cells become None (a genuinely empty cell).
    """
    v = value.strip()
    if not v:
        return None
    # Preserve leading-zero strings as text (007, 02134) — but plain "0" is a number.
    if (
        len(v.lstrip("-")) > 1
        and v.lstrip("-").startswith("0")
        and not v.lstrip("-").startswith("0.")
    ):
        return value
    if _INT_RE.match(v):
        return int(v)
    if _FLOAT_RE.match(v):
        return float(v)
    return value


def _consume_list(doc: "Document", lines: list[str], start: int, ordered: bool) -> int:
    """Consume consecutive list items starting at `start`. Returns next line index."""
    pattern = _NUMBERED_RE if ordered else _BULLET_RE
    style_name = "List Number" if ordered else "List Bullet"
    i = start
    while i < len(lines):
        match = pattern.match(lines[i])
        if not match:
            break
        para = doc.add_paragraph(style=style_name)
        _apply_inline_formatting(para, match.group(1))
        i += 1
    return i


def _consume_table(doc: "Document", lines: list[str], start: int) -> int:
    """Consume a markdown table starting at `start`. Returns next line index."""
    from docx.shared import Pt

    rows_text: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            break
        if _TABLE_SEP_RE.match(line):
            i += 1
            continue
        row_match = _TABLE_ROW_RE.match(line)
        if row_match:
            cells = [c.strip() for c in row_match.group(1).split("|")]
            rows_text.append(cells)
            i += 1
        else:
            break

    if not rows_text:
        return i

    num_cols = max(len(row) for row in rows_text)
    for row in rows_text:
        while len(row) < num_cols:
            row.append("")

    table = doc.add_table(rows=len(rows_text), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows_text):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            if row_idx == 0:
                run = para.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)
            else:
                _apply_inline_formatting(para, cell_text, font_size=Pt(10))

    doc.add_paragraph()
    return i


def _apply_inline_formatting(
    para: "Paragraph", text: str, clear_existing: bool = False, font_size=None
) -> None:
    """Parse inline markdown formatting and add runs to the paragraph."""
    if clear_existing:
        para.clear()

    for match in _INLINE_RE.finditer(text):
        bold_italic, bold, italic, plain = match.groups()
        if bold_italic:
            run = para.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold:
            run = para.add_run(bold)
            run.bold = True
        elif italic:
            run = para.add_run(italic)
            run.italic = True
        elif plain:
            run = para.add_run(plain)
        else:
            continue
        if font_size:
            run.font.size = font_size
